from fpdf import FPDF
from datetime import datetime
import base64
import math
import os
import tempfile
from PIL import Image as _PILImage

# Windows フォントパス（游明朝 → 游ゴシック → メイリオ の優先順）
_FONT_MINCHO_R = r"C:\Windows\Fonts\yumin.ttf"
_FONT_MINCHO_B = r"C:\Windows\Fonts\yumindb.ttf"
_FONT_GOTHIC_R = r"C:\Windows\Fonts\YuGothR.ttc"
_FONT_GOTHIC_B = r"C:\Windows\Fonts\YuGothB.ttc"
_FONT_R = r"C:\Windows\Fonts\meiryo.ttc"
_FONT_B = r"C:\Windows\Fonts\meiryob.ttc"
_FONT_FALLBACK = r"C:\Windows\Fonts\msgothic.ttc"

# Linux フォントパス（Streamlit Cloud / Ubuntu — fonts-noto-cjk パッケージ）
_LINUX_FONT_DIRS = [
    "/usr/share/fonts/opentype/noto",
    "/usr/share/fonts/truetype/noto",
    "/usr/share/fonts/noto",
    "/usr/share/fonts",
]
_LINUX_FONT_R_NAMES = ["NotoSansCJK-Regular.ttc", "NotoSansCJKjp-Regular.otf", "NotoSansJP-Regular.ttf"]
_LINUX_FONT_B_NAMES = ["NotoSansCJK-Bold.ttc", "NotoSansCJKjp-Bold.otf", "NotoSansJP-Bold.ttf"]

# ロゴ・キャラクター画像パス
_LOGO_PATH           = os.path.join(os.path.dirname(__file__), "logo.png")
_CHARACTER_PATH      = os.path.join(os.path.dirname(__file__), "profile_dark.png")
_CHARACTER_WORK_PATH = os.path.join(os.path.dirname(__file__), "profile_work.png")

# 変換済み一時ファイルキャッシュ
_converted_cache: dict = {}

def _make_cover_char(src_path: str, bg=(8, 8, 8), fade_frac: float = 0.45) -> str:
    """表紙用キャラクター: 左・上方向を多方向グラデーションで背景に自然に溶け込ませる。"""
    key = ("cover_char_v3", src_path, bg, fade_frac)
    if key in _converted_cache and os.path.exists(_converted_cache[key]):
        return _converted_cache[key]
    from PIL import ImageDraw
    img = _PILImage.open(src_path).convert("RGBA")
    w, h = img.size
    # アルファマスク（255=完全不透明、0=完全透明）
    mask = _PILImage.new("L", (w, h), 255)

    # 左フェード（45%幅・なだらかな曲線）
    fade_w = int(w * fade_frac)
    for x in range(fade_w):
        t = x / fade_w
        alpha = int(255 * (t ** 1.6))
        for y in range(h):
            cur = mask.getpixel((x, y))
            mask.putpixel((x, y), min(cur, alpha))

    # 上フェード（20%高さ・ホワイトボード上端を自然に溶かす）
    fade_h = int(h * 0.20)
    for y in range(fade_h):
        t = y / fade_h
        alpha = int(255 * (t ** 1.6))
        for x in range(w):
            cur = mask.getpixel((x, y))
            mask.putpixel((x, y), min(cur, alpha))

    # 下フェード（18%高さ・キャラ下端を背景に自然に溶かす）
    fade_h_b = int(h * 0.18)
    for y in range(h - fade_h_b, h):
        t = (h - y) / fade_h_b
        alpha = int(255 * (t ** 1.6))
        for x in range(w):
            cur = mask.getpixel((x, y))
            mask.putpixel((x, y), min(cur, alpha))

    img.putalpha(mask)
    bg_img = _PILImage.new("RGB", (w, h), bg)
    bg_img.paste(img, mask=img.split()[-1])
    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    bg_img.save(tmp.name, "PNG")
    _converted_cache[key] = tmp.name
    return tmp.name

def _to_rgb_path(src_path: str, bg=(8, 8, 8)) -> str:
    """RGBA/P モードの画像を指定背景色でRGBに合成し、一時ファイルパスを返す。"""
    key = (src_path, bg)
    if key in _converted_cache and os.path.exists(_converted_cache[key]):
        return _converted_cache[key]
    img = _PILImage.open(src_path)
    if img.mode in ("RGBA", "LA", "P"):
        background = _PILImage.new("RGB", img.size, bg)
        if img.mode == "P":
            img = img.convert("RGBA")
        background.paste(img, mask=img.split()[-1] if img.mode in ("RGBA", "LA") else None)
        img = background
    elif img.mode != "RGB":
        img = img.convert("RGB")
    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    img.save(tmp.name, "PNG")
    _converted_cache[key] = tmp.name
    return tmp.name

def _logo_on_pdf_bg(src_path: str, target_bg=(8, 8, 8), tolerance=25) -> str:
    """ロゴ画像の背景色をPDF背景色に統一する。コーナーピクセルで背景色を自動検出して置換。"""
    key = ("logo_bg_match", src_path, target_bg, tolerance)
    if key in _converted_cache and os.path.exists(_converted_cache[key]):
        return _converted_cache[key]
    img = _PILImage.open(src_path)
    if img.mode == "P":
        img = img.convert("RGBA")
    if img.mode == "RGBA":
        bg = _PILImage.new("RGB", img.size, target_bg)
        bg.paste(img, mask=img.split()[3])
        img = bg
    elif img.mode != "RGB":
        img = img.convert("RGB")
    w, h = img.size
    arr = img.load()
    sr, sg, sb = arr[0, 0]  # 左上コーナーの色 = 背景色として検出
    for y in range(h):
        for x in range(w):
            r, g, b = arr[x, y]
            if abs(r - sr) <= tolerance and abs(g - sg) <= tolerance and abs(b - sb) <= tolerance:
                arr[x, y] = target_bg
    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    img.save(tmp.name, "PNG")
    _converted_cache[key] = tmp.name
    return tmp.name

# ブランドカラー (R, G, B) — Black & Gold テーマ
NAVY      = (18, 18, 18)      # 本文・テキスト用の深黒（旧NAVY代替）
NAVY_DARK = (8, 8, 8)         # カバー背景用の最深黒
PURPLE    = (107, 70, 193)
WHITE     = (255, 255, 255)
LIGHT     = (248, 246, 240)   # 温かみのあるオフホワイト
GRAY      = (130, 125, 115)   # ウォームグレー
GREEN_BG  = (236, 253, 245)
GREEN_T   = (21, 128, 61)
RED_BG    = (254, 242, 242)
RED_T     = (185, 28, 28)
AMBER_BG  = (255, 251, 220)
AMBER_T   = (146, 64, 14)
BLUE_MID  = (37, 99, 235)
GREEN_MID = (22, 163, 74)
GOLD      = (212, 175, 55)    # メインゴールド
GOLD_LIGHT = (235, 205, 100)  # 明るいゴールド（ハイライト用）
GOLD_BG   = (252, 248, 230)   # ゴールド背景色
CHARCOAL  = (35, 32, 28)      # 深いチャコール（カード背景）
CHARCOAL2 = (50, 46, 40)      # 少し明るいチャコール


def _calc_power_score(site_report: dict) -> tuple:
    """AI総合パワースコア（100点満点）: DRM(30) + BrandZ(40) + GEO(30)。"""
    mi = site_report.get("marketing_insights", {})
    drm = mi.get("drm_score", "C")
    drm_pts = {"A": 30, "B": 22, "C": 14, "D": 6}.get(drm, 14)

    bz = site_report.get("brandz_score", {})
    bz_vals = []
    for key in ("meaningful", "different", "salient"):
        v = bz.get(key, {}).get("score", 5)
        try:
            bz_vals.append(int(v))
        except (ValueError, TypeError):
            bz_vals.append(5)
    bz_avg = sum(bz_vals) / len(bz_vals) if bz_vals else 5.0
    bz_pts = round(bz_avg / 10 * 40)

    geo = site_report.get("geo_score", {})
    try:
        geo_num = int(geo.get("score", 5))
    except (ValueError, TypeError):
        geo_num = 5
    geo_pts = round(geo_num / 10 * 30)

    return drm_pts + bz_pts + geo_pts, drm_pts, bz_pts, geo_pts


def _find_linux_font(names: list) -> str:
    """Linux環境でNoto CJKフォントを探索する。"""
    for d in _LINUX_FONT_DIRS:
        for name in names:
            path = os.path.join(d, name)
            if os.path.exists(path):
                return path
    return ""


def _resolve_font():
    """游明朝 → 游ゴシック → メイリオ → Noto CJK（Linux）の優先順でフォントを解決。"""
    if os.path.exists(_FONT_MINCHO_R) and os.path.exists(_FONT_MINCHO_B):
        return _FONT_MINCHO_R, _FONT_MINCHO_B
    if os.path.exists(_FONT_GOTHIC_R) and os.path.exists(_FONT_GOTHIC_B):
        return _FONT_GOTHIC_R, _FONT_GOTHIC_B
    if os.path.exists(_FONT_R) and os.path.exists(_FONT_B):
        return _FONT_R, _FONT_B
    if os.path.exists(_FONT_FALLBACK):
        return _FONT_FALLBACK, _FONT_FALLBACK
    # Linux (Streamlit Cloud) fallback
    r = _find_linux_font(_LINUX_FONT_R_NAMES)
    b = _find_linux_font(_LINUX_FONT_B_NAMES)
    if r:
        return r, (b if b else r)
    return None, None


class _Base(FPDF):
    def __init__(self):
        super().__init__(orientation="P", unit="mm", format="A4")
        self.alias_nb_pages()
        r_path, b_path = _resolve_font()
        if r_path:
            self.add_font("M", fname=r_path)
            self.add_font("M", style="B", fname=b_path)
            self._font = "M"
        else:
            self._font = "Helvetica"
        self.set_auto_page_break(auto=True, margin=22)
        self.set_margins(15, 32, 15)
        self._accent = PURPLE  # サブクラスで上書き可能

    def header(self):
        self.set_fill_color(*NAVY)
        self.rect(0, 0, 210, 28, style="F")
        self.set_y(6)
        self.set_font(self._font, "B", 13)
        self.set_text_color(*WHITE)
        self.cell(0, 8, "トドク VCT  仮想顧客テストレポート", align="C")
        self.set_y(16)
        self.set_font(self._font, "", 8)
        self.set_text_color(160, 160, 210)
        self.cell(0, 6, "Powered by LIFE DESIGN LAB", align="C")
        self.ln(14)
        self.set_text_color(0, 0, 0)

    def footer(self):
        self.set_y(-14)
        self.set_font(self._font, "", 8)
        self.set_text_color(*GRAY)
        self.cell(0, 10, f"© LIFE DESIGN LAB  ·  {datetime.now().strftime('%Y年%m月%d日')}  ·  {self.page_no()} / {{nb}} ページ", align="C")

    def section_bar(self, text: str):
        self.set_fill_color(*PURPLE)
        self.set_text_color(*WHITE)
        self.set_font(self._font, "B", 9)
        y = self.get_y()
        self.rect(self.l_margin, y, 180, 7.5, style="F")
        self.set_xy(self.l_margin + 3, y + 1)
        self.cell(174, 5.5, text)
        self.ln(10)
        self.set_text_color(0, 0, 0)

    def colored_block(self, text: str, bg, tc, indent=0):
        lm = self.l_margin + indent
        w = 180 - indent
        self.set_fill_color(*bg)
        self.set_text_color(*tc)
        self.set_font(self._font, "", 9)
        self.set_x(lm)
        self.multi_cell(w, 5.5, text, fill=True, padding=(2, 3, 2, 3))
        self.ln(2)
        self.set_text_color(0, 0, 0)

    def meta_line(self, profession, persona_count, test_mode):
        self.set_font(self._font, "", 8)
        self.set_text_color(*GRAY)
        self.set_x(self.l_margin)
        self.cell(0, 5, (
            f"テスト日時: {datetime.now().strftime('%Y年%m月%d日  %H:%M')}"
            f"　｜　ジャンル: {profession}"
            f"　｜　仮想顧客数: {persona_count}人"
            f"　｜　モード: {test_mode}"
        ))
        self.ln(8)
        self.set_text_color(0, 0, 0)

    def metric_box(self, label, value, bg, tc, x, y, w=43, h=20):
        self.set_fill_color(*bg)
        self.rect(x, y, w, h, style="F")
        # 背景が暗い場合はラベルを白っぽく、明るい場合はグレーに
        label_color = (200, 200, 220) if tc == WHITE else GRAY
        self.set_xy(x, y + 2)
        self.set_font(self._font, "", 7)
        self.set_text_color(*label_color)
        self.cell(w, 4, label, align="C")
        self.set_xy(x, y + 7)
        self.set_font(self._font, "B", 13)
        self.set_text_color(*tc)
        self.cell(w, 9, str(value), align="C")
        self.set_text_color(0, 0, 0)

    def cta_footer(self):
        self.ln(6)
        lm = self.l_margin
        y = self.get_y()
        self.set_fill_color(*NAVY_DARK)
        self.rect(lm, y, 180, 18, style="F")
        self.set_fill_color(*GOLD)
        self.rect(lm, y, 180, 0.8, style="F")
        self.rect(lm, y + 17.2, 180, 0.8, style="F")
        self.set_xy(lm, y + 4)
        self.set_font(self._font, "B", 10)
        self.set_text_color(*GOLD)
        self.cell(180, 6, "このレポートはトドクがプロ視点で分析しました", align="C")
        self.set_y(y + 11)
        self.set_font(self._font, "", 8)
        self.set_text_color(160, 148, 100)
        self.cell(180, 5, "LP改善・マンガ制作・コンテンツ戦略のご相談はLIFE DESIGN LABまで", align="C")
        self.set_text_color(0, 0, 0)

    def _final_footer(self):
        """最終ページ専用: ブランドメッセージ + CTAを1本の帯に統合。ロゴ・キャラ付き。"""
        remaining = 297 - self.get_y() - 22
        y = self.get_y() + max(6, remaining * 0.3)
        lm = self.l_margin
        box_h = 28  # 統合帯の高さ

        # 黒背景帯（横幅いっぱい）
        self.set_fill_color(*NAVY_DARK)
        self.rect(0, y, 210, box_h, style="F")
        # 上下ゴールドライン（横幅いっぱい）
        self.set_fill_color(*GOLD)
        self.rect(0, y, 210, 0.8, style="F")
        self.rect(0, y + box_h - 0.8, 210, 0.8, style="F")

        # テキスト（ロゴなし・全幅中央揃え）
        self.set_xy(0, y + 4)
        self.set_font(self._font, "B", 9.5)
        self.set_text_color(*GOLD)
        self.cell(210, 6, "あなたのサイトを 選ばれる存在へ。", align="C")
        self.set_xy(0, y + 12)
        self.set_font(self._font, "", 7.5)
        self.set_text_color(*WHITE)
        self.cell(210, 5, "このレポートはトドクがプロ視点で分析しました", align="C")
        self.set_xy(0, y + 19)
        self.set_font(self._font, "", 6.5)
        self.set_text_color(160, 148, 100)
        self.cell(210, 4.5, "LP改善・マンガ制作・コンテンツ戦略  ·  LIFE DESIGN LAB", align="C")
        self.set_text_color(0, 0, 0)

    def _drm_score_colors(self, mi: dict):
        score = mi.get("drm_score", "C")
        score_color = {"A": GREEN_MID, "B": BLUE_MID, "C": (217, 119, 6), "D": (185, 28, 28)}.get(score, GRAY)
        desc = {
            "A": "集客・教育・販売がすべて機能。広告をかけると費用対効果が高い状態。",
            "B": "基本はできているが1〜2か所の改善で大きく伸びる伸びしろがある状態。",
            "C": "構造に問題あり。改善なく広告をかけると費用を無駄にする可能性が高い。",
            "D": "根本から作り直しが必要。現状のまま運用しても成果は出にくい。",
        }.get(score, "")
        return score, score_color, desc

    def drm_overview(self, mi: dict):
        """p.4用: DRM解説・グレードカード・スコアバッジ・地味だけどヤバいポイント"""
        if not mi:
            return
        lm = self.l_margin
        score, score_color, desc = self._drm_score_colors(mi)

        self.section_bar("マーケティング総合評価（DRM：ダイレクトレスポンスマーケティング）")
        lm = self.l_margin

        # ── DRM 3ステップ フロービジュアル ──
        flow_y = self.get_y()
        flow_steps = [
            ("集客", "認知・アクセス獲得", GREEN_MID),
            ("教育", "信頼・ファン化", BLUE_MID),
            ("販売", "問合せ・成約", GOLD),
        ]
        step_w, arrow_w, flow_h = 50, 10, 24
        total_w = step_w * 3 + arrow_w * 2
        sx_start = lm + (180 - total_w) / 2
        for i, (label, sub, color) in enumerate(flow_steps):
            sx = sx_start + i * (step_w + arrow_w)
            self.set_fill_color(*color)
            self.rect(sx, flow_y, step_w, flow_h, style="F")
            self.set_fill_color(255, 255, 255)
            self.rect(sx, flow_y + flow_h - 1.5, step_w, 1.5, style="F")
            self.set_xy(sx, flow_y + 3.5)
            self.set_font(self._font, "B", 16)
            self.set_text_color(*WHITE)
            self.cell(step_w, 9, label, align="C")
            self.set_xy(sx, flow_y + 13.5)
            self.set_font(self._font, "", 6.5)
            self.set_text_color(*WHITE)
            self.cell(step_w, 5, sub, align="C")
            if i < 2:
                self.set_xy(sx_start + (i + 1) * step_w + i * arrow_w, flow_y + 7)
                self.set_font(self._font, "B", 13)
                self.set_text_color(*GOLD)
                self.cell(arrow_w, 10, ">", align="C")
        self.set_y(flow_y + flow_h + 5)

        # DRM解説ボックス（動的高さ）
        ex_y = self.get_y()
        _drm_explain = (
            "集客・教育・販売の3ステップで見込み客を顧客へと育てるマーケティング手法です。"
            "単に広告を出すのではなく、信頼関係を築きながら自然な流れで問い合わせ・成約へ導く導線設計が核心です。"
        )
        _drm_lines = max(2, -(-len(_drm_explain) // 36))
        ex_h = max(30, _drm_lines * 5 + 16)
        self.set_fill_color(248, 244, 234)
        self.rect(lm, ex_y, 180, ex_h, style="F")
        self.set_fill_color(*GOLD)
        self.rect(lm, ex_y, 3, ex_h, style="F")
        self.set_xy(lm + 6, ex_y + 2.5)
        self.set_font(self._font, "B", 7.5)
        self.set_text_color(*GOLD)
        self.cell(170, 4.5, "DRM（ダイレクトレスポンスマーケティング）とは")
        self.set_xy(lm + 6, ex_y + 8)
        self.set_font(self._font, "", 7.5)
        self.set_text_color(22, 18, 10)
        self.set_auto_page_break(auto=False)
        self.multi_cell(171, 4.5, _drm_explain)
        self.set_auto_page_break(auto=True, margin=22)
        self.set_y(ex_y + ex_h + 4)

        # A〜D グレード説明（横一列）
        grade_y = self.get_y()
        grades = [
            ("A", GREEN_MID,       "広告投下タイミング",
             "集客・教育・販売の導線がすべて機能。今すぐ広告をかけると費用対効果が高い。"),
            ("B", BLUE_MID,        "改善で大きく伸びる",
             "基本はできているが1〜2か所の改善で大幅アップが狙える伸びしろ十分な状態。"),
            ("C", (217, 119, 6),   "改善してから広告を",
             "構造的な問題あり。このまま広告費をかけると無駄になる可能性が高い。"),
            ("D", (185, 28, 28),   "根本的な見直しが必要",
             "集客の仕組みが機能していない。まず導線設計から作り直すことを推奨。"),
        ]
        gw = 43
        for i, (g, gc, g_title, g_desc) in enumerate(grades):
            gx = lm + i * (gw + 2)
            # 背景
            self.set_fill_color(248, 244, 234)
            self.rect(gx, grade_y, gw, 26, style="F")
            # グレードバッジ
            self.set_fill_color(*gc)
            self.rect(gx, grade_y, gw, 8, style="F")
            self.set_xy(gx, grade_y + 0.5)
            self.set_font(self._font, "B", 12)
            self.set_text_color(*WHITE)
            self.cell(gw, 7, g, align="C")
            # タイトル
            self.set_xy(gx + 2, grade_y + 9.5)
            self.set_font(self._font, "B", 6.5)
            self.set_text_color(*gc)
            self.cell(gw - 4, 4, g_title, align="C")
            # 説明文
            self.set_xy(gx + 2, grade_y + 14.5)
            self.set_font(self._font, "", 6)
            self.set_text_color(75, 65, 45)
            self.multi_cell(gw - 4, 3.5, g_desc)
        self.set_y(grade_y + 30)
        self.set_text_color(0, 0, 0)
        self.ln(2)

        # 今回のスコアバッジ + 総評
        y0 = self.get_y()
        self.set_fill_color(42, 36, 22)
        self.rect(lm, y0, 180, 28, style="F")
        self.set_fill_color(*score_color)
        self.rect(lm + 4, y0 + 4, 18, 20, style="F")
        self.set_xy(lm + 4, y0 + 8)
        self.set_font(self._font, "B", 14)
        self.set_text_color(*WHITE)
        self.cell(18, 10, score, align="C")
        self.set_xy(lm + 26, y0 + 4)
        self.set_font(self._font, "B", 8)
        self.set_text_color(160, 148, 100)
        self.cell(150, 5, "今回のサイトの総合評価")
        self.set_xy(lm + 26, y0 + 10)
        self.set_font(self._font, "B", 8.5)
        self.set_text_color(*WHITE)
        self.set_auto_page_break(auto=False)
        self.multi_cell(150, 5.5, desc)
        self.set_auto_page_break(auto=True, margin=22)
        self.set_y(y0 + 32)
        self.set_text_color(0, 0, 0)

        # 地味だけどヤバい改善ポイント
        gem = mi.get("hidden_gem", "")
        if gem:
            if (297 - self.get_y() - 22) < 36:
                self.add_page()
            y_gem = self.get_y()
            # 高さを事前計算（動的）
            _gem_lines = max(2, -(-len(gem) // 40))
            _gem_body_h = _gem_lines * 5.5 + 10
            # ヘッダーバー
            self.set_fill_color(*self._accent)
            self.rect(lm, y_gem, 180, 7, style="F")
            self.set_xy(lm + 5, y_gem + 1)
            self.set_font(self._font, "B", 8.5)
            self.set_text_color(*WHITE)
            self.cell(170, 5, "◆  地味だけどヤバい改善ポイント")
            # 本文エリア（動的高さ）
            self.set_fill_color(250, 244, 222)
            self.rect(lm, y_gem + 7, 180, _gem_body_h, style="F")
            self.set_fill_color(*self._accent)
            self.rect(lm, y_gem + 7, 3, _gem_body_h, style="F")
            self.set_xy(lm + 7, y_gem + 11)
            self.set_font(self._font, "", 9)
            self.set_text_color(*NAVY)
            self.set_auto_page_break(auto=False)
            self.multi_cell(170, 5.5, gem)
            self.set_auto_page_break(auto=True, margin=22)
            self.set_y(y_gem + 7 + _gem_body_h + 4)
            self.ln(2)

        self.set_text_color(0, 0, 0)

    def drm_axis(self, mi: dict):
        """p.5用: DRM 4軸詳細カード"""
        if not mi:
            return
        lm = self.l_margin
        self.section_bar("DRM 4軸詳細分析")
        axis_items = [
            ("ファーストビュー", "最初に目に入る画面・第一印象", "first_view"),
            ("信頼構築", "実績・口コミ・資格で安心感を作る施策", "trust_building"),
            ("差別化", "競合との違いの見せ方・独自ポジション", "differentiation"),
            ("CTA（行動喚起）", "問い合わせや申込みを促すボタン・文章", "cta_strength"),
        ]
        for label, sub, key in axis_items:
            txt = mi.get(key, "")
            if (297 - self.get_y() - 22) < 40:
                self.add_page()
            y = self.get_y()

            # body_h を先に計算してカード全体の高さを確定
            chars_per_line = 55
            lines = max(1, -(-len(txt) // chars_per_line))
            body_h = lines * 5.5 + 7
            card_h = 11.8 + body_h

            # カード全体クリーム背景
            self.set_fill_color(248, 244, 234)
            self.rect(lm, y, 180, card_h, style="F")
            # ゴールド上ライン・左バー（全高）
            self.set_fill_color(*GOLD)
            self.rect(lm, y, 180, 0.8, style="F")
            self.rect(lm, y, 3, card_h, style="F")
            # ヘッダー/ボディ区切り
            self.set_fill_color(201, 169, 110)
            self.rect(lm + 3, y + 11.4, 177, 0.3, style="F")

            # ヘッダーテキスト（ダーク）
            self.set_xy(lm + 7, y + 3)
            self.set_font(self._font, "B", 9)
            self.set_text_color(42, 36, 22)
            label_w = self.get_string_width(label)
            self.cell(label_w + 2, 5.5, label)
            self.set_font(self._font, "", 6.5)
            self.set_text_color(120, 110, 80)
            self.cell(160, 5.5, f"  —  {sub}")

            # 本文テキスト
            self.set_auto_page_break(auto=False)
            self.set_xy(lm + 8, y + 15)
            self.set_font(self._font, "", 9)
            self.set_text_color(22, 18, 10)
            self.multi_cell(168, 6, txt)
            self.set_auto_page_break(auto=True, margin=22)
            self.set_y(y + card_h + 5)
        self.set_text_color(0, 0, 0)

    def drm_section(self, mi: dict):
        """後方互換ラッパー（overview + axis を連続実行）"""
        self.drm_overview(mi)
        self.drm_axis(mi)

    def brandz_section(self, bz: dict):
        """BrandZ（Kantar）3軸スコアページ"""
        if not bz:
            return
        lm = self.l_margin
        self.section_bar("BrandZ スコア（Kantar ブランドエクイティ分析）")

        axes = [
            ("意味性  Meaningful", "顧客ニーズへの機能的・感情的適合度", "meaningful", (180, 140, 60)),
            ("差別性  Different",  "競合との明確な差別化・唯一のポジション", "different",  (100, 140, 80)),
            ("顕著性  Salient",    "想起容易性・AI検索（GEO）での発見可能性", "salient",   (80, 120, 180)),
        ]

        for label, sub, key, accent in axes:
            data = bz.get(key, {})
            score = int(data.get("score", 5)) if str(data.get("score", "5")).isdigit() else 5
            comment = " ".join(data.get("comment", "").split())  # 余分な空白を正規化

            chars_per_line = 52
            lines = max(1, -(-len(comment) // chars_per_line))
            body_h = lines * 5.5 + 7
            card_h = 14 + body_h

            if (297 - self.get_y() - 22) < card_h + 4:
                self.add_page()

            y = self.get_y()
            # カード背景
            self.set_fill_color(248, 244, 234)
            self.rect(lm, y, 180, card_h, style="F")
            self.set_fill_color(*GOLD)
            self.rect(lm, y, 180, 0.8, style="F")  # 太い上線は端から端まで
            self.set_fill_color(*accent)
            self.rect(lm, y, 3, card_h, style="F")

            # スコアバッジ（右側 — 大きめの円 + /10表記）
            cx = lm + 168
            cy = y + card_h / 2
            r = 10
            self.set_fill_color(*accent)
            self.ellipse(cx - r, cy - r, r * 2, r * 2, style="F")
            self.set_font(self._font, "B", 13)
            self.set_text_color(255, 255, 255)
            self.set_xy(cx - r, cy - 5)
            self.cell(r * 2, 8, str(score), align="C")
            self.set_xy(cx - r, cy + 3)
            self.set_font(self._font, "", 6)
            self.cell(r * 2, 4, "/10", align="C")

            # ヘッダー
            self.set_xy(lm + 7, y + 3)
            self.set_font(self._font, "B", 9)
            self.set_text_color(42, 36, 22)
            self.cell(100, 5.5, label)
            self.set_font(self._font, "", 6.5)
            self.set_text_color(120, 110, 80)
            self.set_xy(lm + 7, y + 8.5)
            self.cell(150, 4, sub)

            # 区切り線（スコア円の手前で止める）
            self.set_fill_color(201, 169, 110)
            self.rect(lm + 3, y + 13.5, 153, 0.3, style="F")

            # コメント（スコア円と重ならないよう幅を調整）
            self.set_auto_page_break(auto=False)
            self.set_xy(lm + 8, y + 16)
            self.set_font(self._font, "", 8.5)
            self.set_text_color(22, 18, 10)
            self.multi_cell(148, 5.5, comment)
            self.set_auto_page_break(auto=True, margin=22)
            self.set_y(y + card_h + 4)

        # BrandZ総評
        total_comment = bz.get("brandz_comment", "")
        if total_comment:
            y = self.get_y() + 2
            self.set_fill_color(42, 36, 22)
            self.rect(lm, y, 180, 0.8, style="F")
            self.set_xy(lm, y + 4)
            self.set_font(self._font, "B", 8.5)
            self.set_text_color(*GOLD)
            self.cell(20, 5, "総評")
            self.set_font(self._font, "", 8.5)
            self.set_text_color(22, 18, 10)
            self.multi_cell(160, 5.5, total_comment)

        self.set_text_color(0, 0, 0)

    def _draw_geo_pentagon(self, cx: float, cy: float, r: float, scores: list, labels: list):
        """GEO 5ステップの五角形レーダーチャートを描画。"""
        n = 5
        angles = [-math.pi / 2 + 2 * math.pi * i / n for i in range(n)]

        # グリッドリング（薄い線を5段階）
        for level in range(1, 6):
            pts = [(cx + r * level / 5 * math.cos(a), cy + r * level / 5 * math.sin(a)) for a in angles]
            is_outer = (level == 5)
            self.set_draw_color(*(GOLD if is_outer else (200, 185, 140)))
            self.set_line_width(0.5 if is_outer else 0.15)
            for i in range(n):
                self.line(pts[i][0], pts[i][1], pts[(i + 1) % n][0], pts[(i + 1) % n][1])

        # 軸線（中心から各頂点へ）
        self.set_draw_color(200, 185, 140)
        self.set_line_width(0.15)
        for a in angles:
            self.line(cx, cy, cx + r * math.cos(a), cy + r * math.sin(a))

        # スコアポリゴン（ゴールド塗り）
        safe_scores = []
        for s in scores:
            try:
                safe_scores.append(min(5, max(0, int(s))))
            except (TypeError, ValueError):
                safe_scores.append(0)
        score_pts = [
            (cx + r * (safe_scores[i] / 5) * math.cos(angles[i]),
             cy + r * (safe_scores[i] / 5) * math.sin(angles[i]))
            for i in range(n)
        ]
        self.set_fill_color(*GOLD_LIGHT)
        self.set_draw_color(*GOLD)
        self.set_line_width(0.7)
        self.polygon(score_pts, style="FD")

        # 各頂点ラベルとスコア値（ラベルを近づけてフォント大きめ）
        label_r = r + 9
        for i, (label, a) in enumerate(zip(labels, angles)):
            lx = cx + label_r * math.cos(a)
            ly = cy + label_r * math.sin(a)
            # 上端ラベルは少し上にずらす
            offset_y = -3 if i == 0 else (-1 if i in (1, 4) else 1)
            self.set_font(self._font, "B", 7.5)
            self.set_text_color(60, 50, 30)
            self.set_xy(lx - 20, ly + offset_y - 5)
            self.cell(40, 4.5, label, align="C")
            self.set_font(self._font, "B", 9)
            self.set_text_color(*GOLD)
            self.set_xy(lx - 20, ly + offset_y - 0.5)
            self.cell(40, 5, f"{safe_scores[i]} / 5", align="C")

        self.set_text_color(0, 0, 0)
        self.set_line_width(0.2)

    def geo_section(self, geo: dict):
        """GEO（AI検索最適化）スコアセクション — 総合スコア＋五角形チャート＋チェックリスト"""
        if not geo:
            return
        lm = self.l_margin
        self.ln(4)

        score = int(geo.get("score", 5)) if str(geo.get("score", "5")).isdigit() else 5

        # ── セクション見出し＋スコアバーを統合（1本のダークバー） ──
        y = self.get_y()
        self.set_fill_color(42, 36, 22)
        self.rect(lm, y, 180, 18, style="F")
        self.set_fill_color(*GOLD)
        self.rect(lm, y, 3, 18, style="F")
        # セクション小見出し
        self.set_font(self._font, "", 7)
        self.set_text_color(160, 140, 80)
        self.set_xy(lm + 8, y + 2)
        self.cell(172, 4, "GEO スコア  /  AI検索最適化（Generative Engine Optimization）")
        # スコア行
        self.set_font(self._font, "B", 10)
        self.set_text_color(*GOLD)
        self.set_xy(lm + 8, y + 8)
        self.cell(65, 6, f"総合スコア：{score} / 10")
        self.set_font(self._font, "", 7.5)
        self.set_text_color(220, 200, 150)
        self.cell(107, 6, "AIに推薦・引用されやすさ（Kantar GEO対策基準）")
        self.set_y(y + 21)

        # ── 五角形レーダーチャート（左）＋ステップ詳細（右） ──
        step_scores = geo.get("step_scores", {})
        steps = [
            ("S1 構造化", "step1_structure"),
            ("S2 How to", "step2_howto"),
            ("S3 Q&A",   "step3_qa"),
            ("S4 比較表", "step4_comparison"),
            ("S5 事例",  "step5_case"),
        ]
        if step_scores:
            chart_y = self.get_y()
            chart_area_h = 72  # 五角形エリアの高さ
            chart_area_w = 90  # 左カラム幅

            # 背景
            self.set_fill_color(248, 244, 234)
            self.rect(lm, chart_y, 180, chart_area_h, style="F")
            self.set_fill_color(*GOLD)
            self.rect(lm, chart_y, 180, 0.6, style="F")

            # 五角形チャート（左エリアの中央）
            pentagon_cx = lm + chart_area_w / 2
            pentagon_cy = chart_y + chart_area_h / 2 + 5
            pentagon_r = 19

            score_vals = [step_scores.get(key, 0) for _, key in steps]
            step_labels = [label for label, _ in steps]
            self._draw_geo_pentagon(pentagon_cx, pentagon_cy, pentagon_r, score_vals, step_labels)

            # 右エリア: 5ステップのスコアバー一覧
            rx = lm + chart_area_w + 3
            rw = 180 - chart_area_w - 3
            step_full_labels = ["S1 構造化（数字・ファクト整理）", "S2 How to（方法・手順コンテンツ）", "S3 Q&A（向く人/向かない人）", "S4 比較表（自社他社比較）", "S5 事例（数字入りビフォーアフター）"]
            for i, ((label, key), full_label) in enumerate(zip(steps, step_full_labels)):
                iy = chart_y + 6 + i * 12
                val = step_scores.get(key, 0)
                try:
                    val_num = min(5, max(0, int(val)))
                except (TypeError, ValueError):
                    val_num = 0
                # ラベル
                self.set_font(self._font, "", 6.5)
                self.set_text_color(60, 50, 30)
                self.set_xy(rx, iy)
                self.cell(rw - 14, 4.5, full_label)
                # スコア数値（右端）
                self.set_font(self._font, "B", 9)
                self.set_text_color(*GOLD)
                self.set_xy(rx + rw - 14, iy - 0.5)
                self.cell(14, 5, f"{val_num}/5", align="R")
                # スコアバー
                bar_y = iy + 5.5
                bar_w = rw - 16
                self.set_fill_color(210, 200, 170)
                self.rect(rx, bar_y, bar_w, 3, style="F")
                fill_ratio = val_num / 5
                fill_col = GREEN_MID if fill_ratio >= 0.7 else ((217, 119, 6) if fill_ratio >= 0.4 else (185, 28, 28))
                if fill_ratio > 0:
                    self.set_fill_color(*fill_col)
                    self.rect(rx, bar_y, bar_w * fill_ratio, 3, style="F")

            self.set_text_color(0, 0, 0)
            self.set_y(chart_y + chart_area_h + 4)

        # ── チェックリスト（縦並び：切れ防止） ──
        ok_items = geo.get("checklist_ok", [])
        ng_items = geo.get("checklist_missing", [])
        if ok_items or ng_items:
            y = self.get_y()
            self.set_fill_color(248, 244, 234)
            self.rect(lm, y, 180, 0.6, style="F")
            self.set_y(y + 3)
            # ✅ 揃っている項目
            if ok_items:
                self.set_x(lm + 4)
                self.set_font(self._font, "B", 7.5)
                self.set_text_color(22, 80, 40)
                self.cell(176, 4.5, f"✅ 揃っている項目（{len(ok_items)}）")
                self.ln(5)
                for ok in ok_items[:6]:
                    self.set_x(lm + 8)
                    self.set_font(self._font, "", 8)
                    self.set_text_color(22, 100, 50)
                    self.multi_cell(172, 4.5, f"✅ {ok}")
            self.ln(2)
            # ❌ 不足している項目
            if ng_items:
                self.set_x(lm + 4)
                self.set_font(self._font, "B", 7.5)
                self.set_text_color(140, 20, 20)
                self.cell(176, 4.5, f"❌ 不足している項目（{len(ng_items)}）")
                self.ln(5)
                for ng in ng_items[:6]:
                    self.set_x(lm + 8)
                    self.set_font(self._font, "", 8)
                    self.set_text_color(160, 30, 30)
                    self.multi_cell(172, 4.5, f"❌ {ng}")
            self.ln(3)

        # ── 原因と改善アクション ──
        for items, title, color in [
            (geo.get("issues", []), "AIに無視される原因", (185, 28, 28)),
            (geo.get("actions", []), "最優先 GEO改善アクション", (22, 120, 60)),
        ]:
            if not items:
                continue
            y = self.get_y()
            self.set_fill_color(*color)
            self.rect(lm, y, 180, 6, style="F")
            self.set_font(self._font, "B", 8)
            self.set_text_color(255, 255, 255)
            self.set_xy(lm + 4, y + 1)
            self.cell(180, 4, title)
            self.set_y(y + 8)
            for item in items[:3]:
                self.set_xy(lm + 6, self.get_y())
                self.set_font(self._font, "", 8.5)
                self.set_text_color(22, 18, 10)
                self.cell(4, 5.5, "•")
                self.multi_cell(170, 5.5, str(item))
            self.ln(2)

        self.set_text_color(0, 0, 0)

    def persona_table(self, personas, reactions, mode):
        self.section_bar("仮想顧客ペルソナ一覧")
        lm = self.l_margin
        for i, (p, r) in enumerate(zip(personas, reactions), 1):
            if mode == "sns":
                verdict = r.get("will_stop_scrolling", "no")
                v_label = {"yes": "◎ 停止する", "maybe": "△ 迷う", "no": "✗ スルー"}.get(verdict, verdict)
                impression = r.get("first_impression", "")
            else:
                verdict = r.get("will_inquire", "maybe")
                v_label = {"yes": "◎ 問い合わせする", "maybe": "△ 迷う", "no": "✗ しない"}.get(verdict, verdict)
                impression = r.get("first_impression", "")
            self.set_fill_color(*LIGHT)
            y0 = self.get_y()
            self.rect(lm, y0, 180, 15, style="F")
            self.set_xy(lm + 2, y0 + 2)
            self.set_font(self._font, "B", 8.5)
            self.set_text_color(*NAVY)
            self.cell(130, 5, f"No.{i}  {p['age']}歳 {p['gender']}  —  {p['problem_type']}")
            self.set_font(self._font, "B", 8.5)
            self.set_text_color(*PURPLE)
            self.cell(44, 5, v_label, align="R")
            self.ln(6)
            self.set_x(lm + 4)
            self.set_font(self._font, "", 7.5)
            self.set_text_color(*GRAY)
            self.cell(174, 4, str(impression)[:80])
            self.ln(6)
            self.set_text_color(0, 0, 0)


# ─────────────────────────────────────────────────────────
# 通常テスト (LP / SNS)
# ─────────────────────────────────────────────────────────
def generate_pdf(
    report: dict, personas: list, reactions: list,
    mode: str, profession: str, persona_count: int,
    main_rate: str, test_mode: str = "LP",
) -> bytes:
    pdf = _Base()
    pdf.add_page()
    lm = pdf.l_margin

    pdf.meta_line(profession, persona_count, test_mode)

    # メトリクス
    y0 = pdf.get_y()
    if mode == "sns":
        stop_yes   = sum(1 for r in reactions if r.get("will_stop_scrolling") == "yes")
        stop_maybe = sum(1 for r in reactions if r.get("will_stop_scrolling") == "maybe")
        like_ct    = sum(1 for r in reactions if r.get("will_like") == "yes")
        visit_ct   = sum(1 for r in reactions if r.get("will_visit_profile") == "yes")
        boxes = [
            ("スクロール停止率", report.get("stop_rate", main_rate), PURPLE, WHITE),
            ("停止する",         f"{stop_yes}人",  GREEN_MID,  WHITE),
            ("迷う",             f"{stop_maybe}人", AMBER_BG,  AMBER_T),
            ("いいね率",         report.get("like_rate", "-"), BLUE_MID, WHITE),
            ("プロフ訪問率",     report.get("profile_visit_rate", "-"), (80,40,160), WHITE),
        ]
    else:
        yes_ct   = sum(1 for r in reactions if r.get("will_inquire") == "yes")
        maybe_ct = sum(1 for r in reactions if r.get("will_inquire") == "maybe")
        no_ct    = sum(1 for r in reactions if r.get("will_inquire") == "no")
        boxes = [
            ("推定問い合わせ率", report.get("inquiry_rate", main_rate), PURPLE,     WHITE),
            ("する",             f"{yes_ct}人",                          GREEN_MID,  WHITE),
            ("迷う",             f"{maybe_ct}人",                        AMBER_BG,   AMBER_T),
            ("しない",           f"{no_ct}人",                           (200,50,50), WHITE),
        ]

    gap, bw = 3, 43
    for i, (label, val, bg, tc) in enumerate(boxes):
        pdf.metric_box(label, val, bg, tc, lm + i * (bw + gap), y0)
    pdf.set_y(y0 + 24)
    pdf.ln(4)

    # 総評
    pdf.section_bar("総評")
    pdf.set_font(pdf._font, "", 10)
    pdf.set_x(lm)
    pdf.multi_cell(180, 6, report.get("summary", ""))
    pdf.ln(4)

    # 強み
    pdf.section_bar("効果的だった点")
    for s in report.get("strengths", []):
        text = f"✓  {s.get('point', '')}  —  {s.get('reason', '')}"
        pdf.colored_block(text, GREEN_BG, GREEN_T)
    pdf.ln(2)

    # 弱み
    pdf.section_bar("改善が必要な箇所")
    for w in report.get("weaknesses", []):
        pdf.colored_block(f"✗  {w.get('point', '')}  —  {w.get('reason', '')}", RED_BG, RED_T)
        if w.get("current_text"):
            pdf.colored_block(f"現在: {w['current_text']}", AMBER_BG, AMBER_T, indent=6)
        if w.get("suggestion"):
            pdf.colored_block(f"改善後: {w['suggestion']}", GREEN_BG, GREEN_T, indent=6)
    pdf.ln(2)

    # 最重要改善
    pdf.section_bar("今すぐやるべき最重要改善")
    pdf.set_fill_color(*AMBER_BG)
    pdf.set_draw_color(*AMBER_T)
    pdf.set_text_color(*AMBER_T)
    pdf.set_font(pdf._font, "B", 10)
    pdf.set_x(lm)
    pdf.multi_cell(180, 6, report.get("priority_action", ""), fill=True, border=1, padding=(3, 4, 3, 4))
    pdf.set_text_color(0, 0, 0)
    pdf.ln(4)

    # DRM分析
    mi = report.get("marketing_insights")
    if mi:
        if pdf.get_y() > 180:
            pdf.add_page()
        pdf.drm_section(mi)

    # ペルソナ一覧
    if pdf.get_y() > 220:
        pdf.add_page()
    pdf.persona_table(personas, reactions, mode)

    pdf.cta_footer()
    return bytes(pdf.output())


# ─────────────────────────────────────────────────────────
# A/B テスト専用
# ─────────────────────────────────────────────────────────
def generate_ab_pdf(
    ab_report: dict, personas: list, reactions_a: list, reactions_b: list,
    text_a: str, text_b: str, mode: str, profession: str, persona_count: int,
) -> bytes:
    pdf = _Base()
    pdf.add_page()
    lm = pdf.l_margin

    pdf.meta_line(profession, persona_count, "A/Bテスト")

    # 勝者バナー
    winner = ab_report.get("winner", "?")
    banner_bg = GREEN_MID if winner in ["A", "B"] else (100, 100, 140)
    pdf.set_fill_color(*banner_bg)
    y_b = pdf.get_y()
    pdf.rect(lm, y_b, 180, 14, style="F")
    pdf.set_xy(lm, y_b + 2)
    pdf.set_font(pdf._font, "B", 14)
    pdf.set_text_color(*WHITE)
    label = f"判定：バージョン {winner} が勝利" if winner in ["A", "B"] else "判定：引き分け"
    pdf.cell(180, 10, label, align="C")
    pdf.ln(16)
    pdf.set_text_color(0, 0, 0)

    # スコアボックス
    y0 = pdf.get_y()
    pdf.metric_box("バージョン A スコア", f"{ab_report.get('score_a', '?')} / 100", LIGHT, NAVY, lm,      y0, w=87)
    pdf.metric_box("バージョン B スコア", f"{ab_report.get('score_b', '?')} / 100", LIGHT, NAVY, lm + 93, y0, w=87)
    pdf.set_y(y0 + 24)
    pdf.ln(2)

    # 勝利理由
    pdf.section_bar("勝利理由")
    pdf.set_x(lm)
    pdf.set_font(pdf._font, "", 10)
    pdf.multi_cell(180, 6, ab_report.get("winner_reason", ""))
    pdf.ln(4)

    # 総評
    pdf.section_bar("総評")
    pdf.set_x(lm)
    pdf.multi_cell(180, 6, ab_report.get("summary", ""))
    pdf.ln(4)

    # A/B それぞれの強み（2カラム）
    pdf.section_bar("各バージョンの強み")
    col_w = 87
    y_start = pdf.get_y()
    x_a, x_b = lm, lm + 93

    # A の強み
    pdf.set_fill_color(*LIGHT)
    pdf.set_font(pdf._font, "B", 9)
    pdf.set_text_color(*NAVY)
    pdf.rect(x_a, y_start, col_w, 7, style="F")
    pdf.set_xy(x_a + 2, y_start + 1)
    pdf.cell(col_w - 4, 5, "バージョン A の強み")
    y_a_body = y_start + 8
    pdf.set_text_color(0, 0, 0)
    for point in ab_report.get("a_strengths", []):
        pdf.set_xy(x_a, y_a_body)
        pdf.set_fill_color(*GREEN_BG)
        pdf.set_text_color(*GREEN_T)
        pdf.set_font(pdf._font, "", 8.5)
        pdf.multi_cell(col_w, 5, f"✓  {point}", fill=True, padding=2)
        y_a_body = pdf.get_y() + 1
    y_after_a = y_a_body

    # B の強み（同じ高さから描画）
    pdf.set_fill_color(*LIGHT)
    pdf.set_font(pdf._font, "B", 9)
    pdf.set_text_color(*NAVY)
    pdf.rect(x_b, y_start, col_w, 7, style="F")
    pdf.set_xy(x_b + 2, y_start + 1)
    pdf.cell(col_w - 4, 5, "バージョン B の強み")
    y_b_body = y_start + 8
    pdf.set_text_color(0, 0, 0)
    for point in ab_report.get("b_strengths", []):
        pdf.set_xy(x_b, y_b_body)
        pdf.set_fill_color(*GREEN_BG)
        pdf.set_text_color(*GREEN_T)
        pdf.set_font(pdf._font, "", 8.5)
        pdf.multi_cell(col_w, 5, f"✓  {point}", fill=True, padding=2)
        y_b_body = pdf.get_y() + 1
    y_after_b = y_b_body

    pdf.set_y(max(y_after_a, y_after_b) + 4)
    pdf.set_text_color(0, 0, 0)

    # 良いとこどり
    pdf.section_bar("両方の良いとこどり — 最強バージョンの方向性")
    pdf.colored_block(ab_report.get("best_of_both", ""), AMBER_BG, AMBER_T)
    pdf.ln(2)

    # 最重要改善
    pdf.section_bar("今すぐやるべき最重要改善")
    pdf.set_fill_color(*AMBER_BG)
    pdf.set_draw_color(*AMBER_T)
    pdf.set_text_color(*AMBER_T)
    pdf.set_font(pdf._font, "B", 10)
    pdf.set_x(lm)
    pdf.multi_cell(180, 6, ab_report.get("priority_action", ""), fill=True, border=1, padding=(3, 4, 3, 4))
    pdf.set_text_color(0, 0, 0)
    pdf.ln(4)

    # DRM分析
    ab_mi = ab_report.get("marketing_insights")
    if ab_mi:
        if pdf.get_y() > 180:
            pdf.add_page()
        pdf.drm_section(ab_mi)

    # ペルソナ
    if pdf.get_y() > 210:
        pdf.add_page()
    pdf.persona_table(personas, reactions_a, mode)

    pdf.cta_footer()
    return bytes(pdf.output())


# ─────────────────────────────────────────────────────────
# サイト全体分析 PDF（高級ホテル調デザイン）
# ─────────────────────────────────────────────────────────

class _SitePDF(_Base):
    """Four Seasons / Aman Tokyo スタイルの高級感あるサイト分析PDF。"""

    def __init__(self):
        super().__init__()
        self._accent = GOLD  # DRM説明ボックスなどをゴールドに

    def header(self):
        if self.page_no() == 1:
            return  # カバーページは自前で描画
        # 黒ヘッダー背景
        self.set_fill_color(*NAVY_DARK)
        self.rect(0, 0, 210, 28, style="F")
        # ゴールド極細ライン（ヘッダー下）
        self.set_fill_color(*GOLD)
        self.rect(0, 28, 210, 0.3, style="F")

        # ロゴ（左端に小さく配置）
        if os.path.exists(_LOGO_PATH):
            self.image(_logo_on_pdf_bg(_LOGO_PATH, target_bg=(8,8,8)), x=4, y=2, w=22)

        # テキストを全幅（210mm）中央揃え
        self.set_xy(0, 5)
        self.set_font(self._font, "", 6.5)
        self.set_text_color(*GOLD)
        self.cell(210, 4.5, "LIFE DESIGN LAB  ·  TODOKU", align="C")
        self.set_xy(0, 12)
        self.set_font(self._font, "B", 10.5)
        self.set_text_color(*WHITE)
        self.cell(210, 6, "サイト全体分析レポート", align="C")
        self.ln(18)
        self.set_text_color(0, 0, 0)

    def footer(self):
        if self.page_no() == 1:
            return
        self.set_fill_color(*GOLD)
        self.rect(0, 293, 210, 1.5, style="F")
        self.set_y(-14)
        self.set_font(self._font, "", 7.5)
        self.set_text_color(*GRAY)
        self.cell(
            0, 10,
            f"© LIFE DESIGN LAB  ·  {datetime.now().strftime('%Y年%m月%d日')}"
            f"  ·  {self.page_no()} / {{nb}} ページ",
            align="C",
        )
        self.set_text_color(0, 0, 0)

    def section_bar(self, text: str):
        """黒背景+ゴールドアクセントのセクションヘッダー。"""
        lm = self.l_margin
        y = self.get_y()
        bar_h = 12
        # ヘッダー背景（温かみのあるダーク）
        self.set_fill_color(42, 36, 22)
        self.rect(lm, y, 180, bar_h, style="F")
        # ゴールド上ライン
        self.set_fill_color(*GOLD)
        self.rect(lm, y, 180, 0.8, style="F")
        # ゴールド縦バー（左アクセント）
        self.rect(lm, y, 3.5, bar_h, style="F")
        # セクションタイトル
        self.set_xy(lm + 8, y + 2.5)
        self.set_font(self._font, "B", 9.5)
        self.set_text_color(*WHITE)
        self.cell(168, 7, text)
        # 下ライン
        self.set_fill_color(*GOLD)
        self.rect(lm, y + bar_h, 180, 0.3, style="F")
        self.ln(16)
        self.set_text_color(0, 0, 0)

    def _fill_page_bottom(self):
        """ページ下部の余白処理（現在は空白のまま残す）。"""
        pass

    def next_actions_page(self, site_report: dict):
        """ネクストアクション3選ページ — クライアントが迷わず動けるよう優先行動を大きく提示。"""
        self.add_page()
        lm = self.l_margin
        self.section_bar("これだけやれば変わる — 優先アクション 3 選")

        # アクションデータを収集
        priority_action = site_report.get("priority_action", "")
        weaknesses = site_report.get("weaknesses", [])
        actions = []
        if priority_action:
            actions.append(("最優先アクション", "今週中に実行", priority_action, GREEN_MID))
        for w in weaknesses:
            if len(actions) >= 3:
                break
            sug = w.get("suggestion", "").strip()
            point = w.get("point", "").strip()
            if sug:
                if len(actions) == 1:
                    actions.append((point, "今月中に実行", sug, BLUE_MID))
                else:
                    actions.append((point, "来月以降に対応", sug, (140, 108, 40)))
        while len(actions) < 3:
            actions.append(("継続的な改善", "随時", "スコアの低い項目から順に対策を進めてください。", (140, 120, 80)))

        numbers = ["1", "2", "3"]
        timing_colors = [GREEN_MID, BLUE_MID, (140, 108, 40)]
        card_h = 50
        gap = 5

        self.set_auto_page_break(auto=False)
        for idx, (title, timing, action, color) in enumerate(actions[:3]):
            y_c = self.get_y()
            # カード全体背景
            self.set_fill_color(248, 244, 234)
            self.rect(lm, y_c, 180, card_h, style="F")
            # 左カラー帯（20mm）
            self.set_fill_color(*color)
            self.rect(lm, y_c, 20, card_h, style="F")
            # 番号
            self.set_xy(lm, y_c + 13)
            self.set_font(self._font, "B", 22)
            self.set_text_color(*WHITE)
            self.cell(20, 18, numbers[idx], align="C")
            # タイミングバッジ
            t_col = timing_colors[idx]
            self.set_fill_color(*t_col)
            self.rect(lm + 24, y_c + 4, 28, 7, style="F")
            self.set_xy(lm + 24, y_c + 4.5)
            self.set_font(self._font, "B", 7)
            self.set_text_color(*WHITE)
            self.cell(28, 6, timing, align="C")
            # アクションタイトル
            self.set_xy(lm + 56, y_c + 4)
            self.set_font(self._font, "B", 7)
            self.set_text_color(120, 100, 50)
            self.cell(120, 7, title)
            # アクション本文（大きめフォント）
            self.set_xy(lm + 24, y_c + 15)
            self.set_font(self._font, "B", 10)
            self.set_text_color(22, 18, 10)
            self.multi_cell(152, 6.5, action)
            self.set_y(y_c + card_h + gap)

        # 締めの一言（バーなし・大きなゴールドテキストのみ）
        self.set_y(self.get_y() + 10)
        self.ln(8)
        self.set_font(self._font, "B", 12)
        self.set_text_color(*GOLD)
        self.cell(180, 9, "まず①だけ実行すれば、問い合わせ数は変わり始めます。", align="C")

        self.set_text_color(0, 0, 0)
        self._final_footer()
        self.set_auto_page_break(auto=True, margin=22)

    def guide_page(self):
        """各スコア・指標の見方を解説するガイドページ（1ページ収録・大学生でもわかる言葉）。"""
        self.add_page()
        lm = self.l_margin
        self.section_bar("このレポートの見方 — 各指標・スコアガイド")

        # ── 見方ガイド（大学生でもわかる簡潔な説明）──
        # タイトル(太字) / ひと言解説 / 判断基準
        guides = [
            (
                "推定問い合わせ率",
                "AIが演じる10人の仮想顧客が「問い合わせしたい」と思った割合です。",
                "30%以上=高水準　12〜29%=標準　11%以下=要改善",
                (22, 120, 60),
            ),
            (
                "AI 総合パワースコア（100点満点）",
                "DRM・BrandZ・GEOの3つのスコアを合計した、サイト全体の総合点です。",
                "80点以上=広告を出すタイミング　60〜79点=改善で伸びる　59点以下=要見直し",
                (180, 140, 20),
            ),
            (
                "DRMスコア（A〜D）",
                "「知る→信頼する→申し込む」という流れがサイトで自然に設計されているか、の評価です。",
                "A=完璧な状態　B=あと一歩　C=構造に問題あり　D=根本から作り直しが必要",
                BLUE_MID,
            ),
            (
                "BrandZ 3軸スコア（各10点）",
                "ブランドの強さを「役に立つか」「他と違うか」「思い出されるか」の3点で採点します。",
                "意味性=悩みに応える　差別性=他社と違う　顕著性=真っ先に思い浮かぶ　3つ揃って強いブランドに",
                (100, 140, 80),
            ),
            (
                "GEO総合スコア（10点満点）",
                "ChatGPTやGeminiなどのAIに「おすすめを教えて」と聞かれたとき、紹介されやすいかの評価です。",
                "7点以上=AIに紹介されやすい　4〜6点=普通　3点以下=AIに無視されている状態",
                (80, 120, 180),
            ),
            (
                "GEO 5ステップ（各5点）",
                "AI対策で必要な5つの要素を採点します。S1=数字で説明　S2=やり方コンテンツ　S3=Q&A　S4=比較表　S5=体験談",
                "点数が低い項目から順に対策するとAIに評価されやすくなります",
                (80, 120, 180),
            ),
            (
                "11項目チェックリスト（AIが求める情報）",
                "AIが「このサイトは信頼できる」と判断するために必要な情報が揃っているかのチェックです。",
                "①誰向け ②悩み ③他社との差 ④料金 ⑤対応エリア ⑥利用の流れ ⑦Q&A ⑧向く人 ⑨向かない人 ⑩口コミ ⑪実績 ⑫専門資格",
                (80, 120, 180),
            ),
        ]

        # 固定高さカード（7枚 × 26mm = 182mm + ヘッダー約28mm = 210mm → 1ページに収まる）
        card_h = 25
        gap = 1
        # 自動改ページを完全に無効化（_final_footer含めて1ページ内で完結させる）
        self.set_auto_page_break(auto=False)

        for title, body, criteria, accent in guides:
            y = self.get_y()
            # カード背景
            self.set_fill_color(248, 244, 234)
            self.rect(lm, y, 180, card_h, style="F")
            self.set_fill_color(*accent)
            self.rect(lm, y, 3, card_h, style="F")
            self.set_fill_color(*GOLD)
            self.rect(lm, y, 155, 0.4, style="F")
            # タイトル
            self.set_xy(lm + 7, y + 2)
            self.set_font(self._font, "B", 8)
            self.set_text_color(42, 36, 22)
            self.cell(165, 5, title)
            # 区切り線
            self.set_fill_color(201, 169, 110)
            self.rect(lm + 3, y + 8.5, 174, 0.25, style="F")
            # ひと言解説
            self.set_xy(lm + 7, y + 10)
            self.set_font(self._font, "", 7)
            self.set_text_color(22, 18, 10)
            self.multi_cell(170, 3.8, body)
            # 判断基準（細字・グレー）
            self.set_xy(lm + 7, y + 18.5)
            self.set_font(self._font, "", 6.5)
            self.set_text_color(100, 90, 60)
            self.multi_cell(170, 3.8, criteria)
            self.set_y(y + card_h + gap)

        self.set_text_color(0, 0, 0)
        # _final_footer は auto=False のまま描画（余分ページを生まない）
        self._final_footer()
        self.set_auto_page_break(auto=True, margin=22)

    def cover_page(self, site_url: str, profession: str, page_count: int, power_score: int = 0):
        self.set_auto_page_break(auto=False)
        self.add_page()

        # ━━ ① 黒背景（最下レイヤー）━━
        self.set_fill_color(*NAVY_DARK)
        self.rect(0, 0, 210, 297, style="F")

        # ━━ ② テキスト（全幅中央揃え） ━━  ※キャラより先に描画
        # ロゴを上に詰めてURLカードがキャラ画像と被らないようにする
        logo_y = 12
        if os.path.exists(_LOGO_PATH):
            logo_size = 36
            self.image(_logo_on_pdf_bg(_LOGO_PATH, target_bg=(8,8,8)), x=105 - logo_size / 2, y=logo_y, w=logo_size)
            logo_y += logo_size + 1
        else:
            self.set_y(logo_y + 4)
            self.set_font(self._font, "B", 12)
            self.set_text_color(*GOLD)
            self.cell(210, 8, "LIFE DESIGN LAB", align="C")
            logo_y += 20

        # ── タイトルブロック（ロゴ下48mmとURLカード96mmの中間に明示配置） ──
        _ty = 57  # タイトルテキストの上端 y
        self.set_fill_color(*GOLD)
        self.rect(105 - 26, _ty - 5, 52, 0.2, style="F")   # 上区切り線
        self.set_xy(0, _ty)
        self.set_font(self._font, "B", 20)
        self.set_text_color(*WHITE)
        self.cell(210, 11, "サイト全体分析レポート", align="C")
        self.set_fill_color(*GOLD)
        self.rect(105 - 26, _ty + 14, 52, 0.2, style="F")  # 下区切り線
        self.set_xy(0, _ty + 20)
        self.set_font(self._font, "", 7)
        self.set_text_color(155, 143, 96)
        self.cell(210, 4, "仮想顧客テスト  ·  Powered by Claude AI", align="C")

        # URLカード（固定位置 — ロゴ下/タイトル/キャラのバランスを固定）
        self.set_y(96)
        bx, by = 14, self.get_y()
        bw, bh = 182, 34
        url_w, score_w = 115, 67

        # カード全体を同一背景色で描画（色の境界線をなくす）
        _card_bg = (14, 12, 8)
        self.set_fill_color(*_card_bg)
        self.rect(bx, by, bw, bh, style="F")
        # 上下ゴールドライン
        self.set_fill_color(*GOLD)
        self.rect(bx, by, bw, 0.3, style="F")
        self.rect(bx, by + bh - 0.3, bw, 0.3, style="F")
        # 左アクセントライン
        self.rect(bx, by, 1.5, bh, style="F")
        # 縦区切り線（URL/スコア境界）
        sx = bx + url_w
        self.rect(sx, by + 4, 0.3, bh - 8, style="F")

        # URL側（左）
        self.set_xy(bx + 6, by + 5)
        self.set_font(self._font, "", 6)
        self.set_text_color(*GOLD)
        self.cell(url_w - 8, 4, "分析対象サイト")
        self.set_xy(bx + 6, by + 11)
        self.set_font(self._font, "B", 6.5)
        self.set_text_color(*WHITE)
        url_disp = site_url[:46] + ("…" if len(site_url) > 46 else "")
        self.cell(url_w - 8, 5.5, url_disp)
        self.set_xy(bx + 6, by + 22)
        self.set_font(self._font, "", 6.5)
        self.set_text_color(155, 143, 96)
        self.cell(url_w - 8, 5, f"ジャンル: {profession}  ·  {page_count}ページ収集")

        # スコア側（右）
        ps_col = GREEN_MID if power_score >= 70 else ((217, 119, 6) if power_score >= 50 else (185, 28, 28))
        ps_str = "広告投下タイミング" if power_score >= 70 else ("改善で伸びる" if power_score >= 50 else "要見直し")
        sx2 = sx + 2
        self.set_xy(sx2, by + 4)
        self.set_font(self._font, "", 6)
        self.set_text_color(155, 143, 96)
        self.cell(score_w - 2, 4, "AI 総合パワースコア", align="C")
        self.set_xy(sx2, by + 9)
        self.set_font(self._font, "B", 24)
        self.set_text_color(*ps_col)
        self.cell(score_w - 12, 12, f"{power_score}", align="R")
        self.set_font(self._font, "", 9)
        self.set_text_color(155, 143, 96)
        self.cell(10, 12, "/100")
        self.set_xy(sx2, by + 26)
        self.set_font(self._font, "B", 6.5)
        self.set_text_color(*ps_col)
        self.cell(score_w - 2, 5, ps_str, align="C")

        # 作成日（ページ幅基準で真ん中）
        self.set_y(by + bh + 9)
        self.set_x(0)
        self.set_font(self._font, "", 7)
        self.set_text_color(115, 105, 70)
        self.cell(210, 5, f"作成日:  {datetime.now().strftime('%Y年%m月%d日')}", align="C")

        # ━━ ③ キャラクター（URLカードの上・金枠の下） ━━
        if os.path.exists(_CHARACTER_PATH):
            try:
                char_path = _make_cover_char(_CHARACTER_PATH, bg=(8, 8, 8), fade_frac=0.45)
                img = _PILImage.open(char_path)
                char_h = 148  # ホワイトボードが見える高さ（小さめで下端が切れない）
                char_w = round(char_h * img.width / img.height, 1)
                self.image(char_path, x=210 - char_w, y=297 - char_h, w=char_w, h=char_h)
            except Exception:
                pass

        # ━━ ④ ゴールドフレーム（最上レイヤー・キャラの上に重ねる） ━━
        self.set_draw_color(*GOLD)
        self.set_line_width(0.4)
        self.rect(4, 4, 202, 289)
        for cx, cy in [(4, 4), (206, 4), (4, 293), (206, 293)]:
            self.set_xy(cx - 1.5, cy - 1.5)
            self.set_font(self._font, "", 4.5)
            self.set_text_color(*GOLD)
            self.cell(3, 3, "●", align="C")

        # 注意書き
        self.set_xy(0, 287)
        self.set_font(self._font, "", 6)
        self.set_text_color(88, 80, 54)
        self.cell(210, 4, "本レポートはトドク（LIFE DESIGN LAB）が作成したものです。", align="C")


def generate_site_pdf(
    site_report: dict,
    scraped_pages: list,
    profession: str,
    device_label: str = "💻 パソコン",
    site_url: str = "",
) -> bytes:
    """サイト全体分析レポートのPDF（高級ホテル調デザイン）。"""
    pdf = _SitePDF()
    lm = pdf.l_margin

    total_ps_cover, _, _, _ = _calc_power_score(site_report)
    pdf.cover_page(site_url, profession, len(scraped_pages), power_score=total_ps_cover)
    pdf.set_auto_page_break(auto=True, margin=22)  # カバー後に再有効化
    pdf.add_page()

    # ── メタ情報行（p.2〜共通） ──
    def _meta_bar(p):
        device_str = device_label.replace("💻 ", "PC").replace("📱 ", "スマホ")
        p.set_font(p._font, "", 7)
        p.set_text_color(*GRAY)
        p.set_x(lm)
        p.cell(0, 5, (
            f"分析日時: {datetime.now().strftime('%Y年%m月%d日  %H:%M')}"
            f"  ｜  ジャンル: {profession}"
            f"  ｜  収集ページ数: {len(scraped_pages)}ページ"
            f"  ｜  閲覧デバイス: {device_str}"
        ))
        p.ln(7)
        p.set_text_color(0, 0, 0)

    _meta_bar(pdf)

    # ─────────────────────────────────────────────
    # p.2  KPI ダッシュボード + 強み/弱み
    # ─────────────────────────────────────────────
    inquiry_rate = site_report.get("inquiry_rate", "—")
    overall      = site_report.get("overall_impression", "")

    try:
        rate_num = int(''.join(filter(str.isdigit, str(inquiry_rate))))
    except ValueError:
        rate_num = 0
    if rate_num >= 30:
        rate_color, rate_badge = GREEN_MID, "高水準"
    elif rate_num >= 12:
        rate_color, rate_badge = (217, 119, 6), "標準"
    else:
        rate_color, rate_badge = (185, 28, 28), "要改善"

    # ── KPI ヒーロー ──
    y0 = pdf.get_y()
    hero_h = 38
    pdf.set_fill_color(*rate_color)
    pdf.rect(lm, y0, 180, 1.5, style="F")
    # 左: 第一印象
    left_w = 112
    pdf.set_fill_color(244, 241, 233)
    pdf.rect(lm, y0 + 1.5, left_w, hero_h - 1.5, style="F")
    pdf.set_fill_color(*GOLD)
    pdf.rect(lm, y0 + 1.5, 3, hero_h - 1.5, style="F")
    pdf.set_xy(lm + 8, y0 + 6)
    pdf.set_font(pdf._font, "", 7)
    pdf.set_text_color(140, 128, 90)
    pdf.cell(100, 4.5, "サイト全体の第一印象")
    pdf.set_xy(lm + 8, y0 + 12)
    pdf.set_font(pdf._font, "", 9)
    pdf.set_text_color(22, 18, 10)
    pdf.multi_cell(100, 5.5, overall[:120])
    # 右: KPI
    rx, rw = lm + left_w + 4, 64
    pdf.set_fill_color(42, 36, 22)
    pdf.rect(rx, y0 + 1.5, rw, hero_h - 1.5, style="F")
    pdf.set_xy(rx, y0 + 6)
    pdf.set_font(pdf._font, "", 7)
    pdf.set_text_color(160, 148, 100)
    pdf.cell(rw, 4.5, "推定問い合わせ率", align="C")
    pdf.set_xy(rx, y0 + 12)
    pdf.set_font(pdf._font, "B", 26)
    pdf.set_text_color(*rate_color)
    pdf.cell(rw, 13, str(inquiry_rate), align="C")
    gy = y0 + 27
    bw_g = rw - 16
    fill_w = min(rate_num / 40.0, 1.0) * bw_g
    pdf.set_fill_color(35, 32, 25)
    pdf.rect(rx + 8, gy, bw_g, 3, style="F")
    if fill_w > 0:
        pdf.set_fill_color(*rate_color)
        pdf.rect(rx + 8, gy, fill_w, 3, style="F")
    pdf.set_xy(rx, gy + 4.5)
    pdf.set_font(pdf._font, "B", 7)
    pdf.set_text_color(*rate_color)
    pdf.cell(rw, 4, rate_badge, align="C")

    pdf.set_y(y0 + hero_h + 6)
    pdf.set_text_color(0, 0, 0)

    # ── AI 総合パワースコア（大型ビジュアル）──
    total_ps, drm_pts, bz_pts, geo_pts = _calc_power_score(site_report)
    ps_color = GREEN_MID if total_ps >= 70 else ((217, 119, 6) if total_ps >= 50 else (185, 28, 28))
    ps_label = "広告投下タイミング" if total_ps >= 70 else ("改善で大きく伸びる" if total_ps >= 50 else "要見直し")
    y_ps = pdf.get_y()
    ps_bar_h = 30
    pdf.set_fill_color(42, 36, 22)
    pdf.rect(lm, y_ps, 180, ps_bar_h, style="F")
    pdf.set_fill_color(*ps_color)
    pdf.rect(lm, y_ps, 180, 1.5, style="F")
    # 大きなスコア数値（左エリア）
    pdf.set_xy(lm + 4, y_ps + 3)
    pdf.set_font(pdf._font, "", 7)
    pdf.set_text_color(160, 148, 100)
    pdf.cell(70, 4, "AI 総合パワースコア")
    pdf.set_xy(lm + 4, y_ps + 8)
    pdf.set_font(pdf._font, "B", 22)
    pdf.set_text_color(*ps_color)
    pdf.cell(26, 13, f"{total_ps}")
    pdf.set_font(pdf._font, "", 9)
    pdf.set_text_color(160, 148, 100)
    pdf.cell(14, 13, "/ 100")
    # ステータスラベル
    pdf.set_xy(lm + 4, y_ps + ps_bar_h - 7)
    pdf.set_font(pdf._font, "B", 7)
    pdf.set_text_color(*ps_color)
    pdf.cell(70, 4.5, f"▶ {ps_label}")
    # プログレスバー（スコア／100）
    pb_x, pb_y, pb_w, pb_h = lm + 4, y_ps + ps_bar_h - 2.5, 68, 2
    pdf.set_fill_color(60, 52, 35)
    pdf.rect(pb_x, pb_y, pb_w, pb_h, style="F")
    pdf.set_fill_color(*ps_color)
    pdf.rect(pb_x, pb_y, pb_w * total_ps / 100, pb_h, style="F")
    # 区切り縦線
    pdf.set_fill_color(70, 62, 45)
    pdf.rect(lm + 82, y_ps + 4, 0.3, ps_bar_h - 8, style="F")
    # サブスコア（右3列）
    x_sub = lm + 86
    for sub_l, sub_v, sub_max in [("DRM", drm_pts, 30), ("BrandZ", bz_pts, 40), ("GEO", geo_pts, 30)]:
        pdf.set_xy(x_sub, y_ps + 4)
        pdf.set_font(pdf._font, "", 6.5)
        pdf.set_text_color(140, 128, 90)
        pdf.cell(30, 4, sub_l, align="C")
        pdf.set_xy(x_sub, y_ps + 9)
        pdf.set_font(pdf._font, "B", 12)
        pdf.set_text_color(*GOLD)
        pdf.cell(30, 7, f"{sub_v}", align="C")
        pdf.set_xy(x_sub, y_ps + 17)
        pdf.set_font(pdf._font, "", 6)
        pdf.set_text_color(100, 90, 60)
        pdf.cell(30, 4, f"/ {sub_max}点", align="C")
        # ミニバー
        mb_x = x_sub + 3
        mb_w = 24
        pdf.set_fill_color(60, 52, 35)
        pdf.rect(mb_x, y_ps + 22, mb_w, 1.5, style="F")
        pdf.set_fill_color(*GOLD)
        pdf.rect(mb_x, y_ps + 22, mb_w * sub_v / sub_max, 1.5, style="F")
        x_sub += 32
    pdf.set_y(y_ps + ps_bar_h + 4)
    pdf.set_text_color(0, 0, 0)

    # ── 強み / 弱み ──
    pdf.section_bar("強み・弱み分析")
    col_w = 87
    y_sw = pdf.get_y()

    strengths = site_report.get("strengths", [])
    pdf.set_fill_color(232, 250, 240)
    pdf.set_font(pdf._font, "B", 8.5)
    pdf.set_text_color(*GREEN_T)
    pdf.rect(lm, y_sw, col_w, 8, style="F")
    pdf.set_xy(lm + 5, y_sw + 2)
    pdf.cell(col_w - 8, 4.5, "◆  強み")
    y_s = y_sw + 10
    for s in strengths:
        txt = f"◆ {s.get('point', '')}\n   {s.get('reason', '')}"
        pdf.set_xy(lm, y_s)
        pdf.set_fill_color(240, 253, 247)
        pdf.set_text_color(*GREEN_T)
        pdf.set_font(pdf._font, "", 8)
        pdf.multi_cell(col_w, 5, txt, fill=True, padding=(3, 4, 3, 4))
        y_s = pdf.get_y() + 3
    y_after_s = y_s

    weaknesses = site_report.get("weaknesses", [])
    wx = lm + col_w + 6
    pdf.set_fill_color(255, 236, 232)
    pdf.set_font(pdf._font, "B", 8.5)
    pdf.set_text_color(*RED_T)
    pdf.rect(wx, y_sw, col_w, 8, style="F")
    pdf.set_xy(wx + 5, y_sw + 2)
    pdf.cell(col_w - 8, 4.5, "▲  弱み・改善点")
    y_w = y_sw + 10
    _priority_labels = ["最優先", "優先対応", "改善推奨"]
    _priority_colors = [(185, 28, 28), (217, 119, 6), (100, 100, 80)]
    for i, w in enumerate(weaknesses):
        p_idx = min(i, 2)
        p_label = _priority_labels[p_idx]
        p_color = _priority_colors[p_idx]
        point  = w.get('point', '').strip()
        reason = w.get('reason', '').strip()
        y_card = y_w
        # タイトル折り返し行数を事前推定（3mm/char @ 7pt、超保守的）
        _pt_cpl = max(1, (col_w - 34) // 3)
        _pt_lines = max(1, -(-len(point) // _pt_cpl))
        header_h = max(7, 2 + _pt_lines * 4.5 + 1.5)
        # バッジ（色付き・動的高さ）
        pdf.set_fill_color(*p_color)
        pdf.rect(wx, y_card, 26, header_h, style="F")
        # タイトル背景
        pdf.set_fill_color(255, 232, 230)
        pdf.rect(wx + 26, y_card, col_w - 26, header_h, style="F")
        # バッジテキスト（縦中央）
        pdf.set_xy(wx, y_card + (header_h - 5) / 2)
        pdf.set_font(pdf._font, "B", 6.5)
        pdf.set_text_color(*WHITE)
        pdf.cell(26, 5, p_label, align="C")
        # タイトルテキスト（折り返し対応・▲はfpdf2のword-break問題で削除）
        pdf.set_xy(wx + 28, y_card + 1.5)
        pdf.set_font(pdf._font, "B", 7)
        pdf.set_text_color(*RED_T)
        pdf.multi_cell(col_w - 30, 4.5, point, padding=(0, 2, 0, 2))
        # ボディ行: 理由テキスト
        pdf.set_xy(wx, y_card + header_h)
        pdf.set_fill_color(255, 245, 243)
        pdf.set_text_color(140, 40, 40)
        pdf.set_font(pdf._font, "", 7.5)
        pdf.multi_cell(col_w, 4.5, f"  {reason}", fill=True, padding=(2, 4, 3, 4))
        y_w = pdf.get_y() + 3
    y_after_w = y_w

    pdf.set_y(max(y_after_s, y_after_w) + 6)
    pdf.set_text_color(0, 0, 0)
    pdf._fill_page_bottom()

    # ─────────────────────────────────────────────
    # p.3  導線問題 + あると効果的なページ
    # ─────────────────────────────────────────────
    pdf.add_page()
    _meta_bar(pdf)

    nav_issues = site_report.get("navigation_issues", [])
    if nav_issues:
        pdf.section_bar("導線・ナビゲーションの問題")
        for ni in nav_issues:
            issue_text = ni.get("issue", "").strip()        # 末尾\nを除去（空行による空白防止）
            suggestion_text = ni.get("suggestion", "").strip()
            page_name = ni.get("page", "").strip()
            # カード高さを超保守的に推定（20文字/行）→ 事前に改ページしてカード内では絶対に改ページしない
            _i_lines = max(1, -(-len(issue_text) // 20))
            _s_lines = max(1, -(-len(suggestion_text) // 22))
            _card_h = 7 + _i_lines * 5.5 + 10 + _s_lines * 5.5 + 10
            if (297 - pdf.get_y() - 22) < _card_h + 6:
                pdf.add_page()
                _meta_bar(pdf)
            y_ni = pdf.get_y()
            # ページ名バッジ（全幅）
            pdf.set_fill_color(42, 36, 22)
            pdf.rect(lm, y_ni, 180, 7, style="F")
            pdf.set_fill_color(*GOLD)
            pdf.rect(lm, y_ni, 3.5, 7, style="F")
            pdf.set_xy(lm + 8, y_ni + 1.5)
            pdf.set_font(pdf._font, "B", 8)
            pdf.set_text_color(*GOLD)
            pdf.cell(168, 4, page_name[:35])
            # 問題テキスト
            pdf.set_xy(lm, y_ni + 8)
            pdf.set_font(pdf._font, "B", 8.5)
            pdf.set_text_color(*RED_T)
            pdf.set_fill_color(255, 245, 245)
            pdf.multi_cell(180, 5.5, issue_text, fill=True, padding=(2, 4, 2, 4))
            # 改善提案ラベル行（細い帯）
            y_sug = pdf.get_y() + 2
            pdf.set_fill_color(220, 205, 150)
            pdf.rect(lm, y_sug, 180, 5, style="F")
            pdf.set_xy(lm + 4, y_sug + 0.5)
            pdf.set_font(pdf._font, "B", 6.5)
            pdf.set_text_color(60, 40, 10)
            pdf.cell(172, 4, "改善提案")
            # 改善提案本文（プレフィックスなし・全幅）
            pdf.set_xy(lm, y_sug + 5)
            pdf.set_fill_color(252, 247, 232)
            pdf.set_font(pdf._font, "", 8)
            pdf.set_text_color(22, 18, 10)
            pdf.multi_cell(180, 5.5, suggestion_text, fill=True, padding=(2, 6, 2, 6))
            pdf.set_y(pdf.get_y() + 5)
        pdf.set_text_color(0, 0, 0)
        pdf.ln(4)

    missing = site_report.get("missing_pages", [])
    if missing:
        if (297 - pdf.get_y() - 22) < 40:
            pdf.add_page()
            _meta_bar(pdf)
        pdf.section_bar("あると効果的なページ（現在なし）")
        for m in missing[:8]:
            _m_lines = max(1, -(-len(m) // 46))
            _m_h = max(9, _m_lines * 5.5 + 4)
            if (297 - pdf.get_y() - 22) < _m_h + 3:
                pdf.add_page()
                _meta_bar(pdf)
            y_m = pdf.get_y()
            pdf.set_fill_color(248, 244, 234)
            pdf.rect(lm, y_m, 180, _m_h, style="F")
            pdf.set_fill_color(*GOLD)
            pdf.rect(lm, y_m, 3, _m_h, style="F")
            pdf.set_xy(lm + 8, y_m + 2)
            pdf.set_font(pdf._font, "B", 9)
            pdf.set_text_color(22, 18, 10)
            pdf.multi_cell(168, 5.5, f"◆  {m}")
            pdf.set_y(y_m + _m_h + 1)
        pdf.set_text_color(0, 0, 0)

    pdf._fill_page_bottom()

    # ─────────────────────────────────────────────
    # p.4  最重要改善 + DRM概要
    # ─────────────────────────────────────────────
    pdf.add_page()
    _meta_bar(pdf)

    pdf.section_bar("今すぐやるべき最重要改善アクション")
    priority = site_report.get("priority_action", "")
    y_p = pdf.get_y()
    box_h = max(22, -(-len(priority) // 42) * 6 + 10)
    pdf.set_fill_color(248, 244, 234)
    pdf.rect(lm, y_p, 180, box_h, style="F")
    pdf.set_fill_color(*GOLD)
    pdf.rect(lm, y_p, 4, box_h, style="F")
    pdf.rect(lm, y_p, 180, 0.5, style="F")
    pdf.rect(lm, y_p + box_h - 0.5, 180, 0.5, style="F")
    pdf.set_xy(lm + 9, y_p + 5)
    pdf.set_font(pdf._font, "B", 9.5)
    pdf.set_text_color(42, 36, 22)
    pdf.multi_cell(167, 6.5, priority)
    pdf.set_y(y_p + box_h + 8)
    pdf.set_text_color(0, 0, 0)

    mi = site_report.get("marketing_insights")
    if mi:
        pdf.drm_overview(mi)

    pdf._fill_page_bottom()

    # ─────────────────────────────────────────────
    # p.5  DRM 4軸詳細
    # ─────────────────────────────────────────────
    if mi:
        pdf.add_page()
        _meta_bar(pdf)
        pdf.drm_axis(mi)
        pdf._fill_page_bottom()

    # ─────────────────────────────────────────────
    # p.6  BrandZ スコア
    # ─────────────────────────────────────────────
    bz = site_report.get("brandz_score")
    if bz:
        pdf.add_page()
        _meta_bar(pdf)
        pdf.brandz_section(bz)
        pdf._fill_page_bottom()

    # ─────────────────────────────────────────────
    # p.7  GEO スコア
    # ─────────────────────────────────────────────
    geo = site_report.get("geo_score")
    if geo:
        pdf.add_page()
        _meta_bar(pdf)
        pdf.geo_section(geo)
        pdf._fill_page_bottom()

    # ─────────────────────────────────────────────
    # p.8  収集ページ一覧
    # ─────────────────────────────────────────────
    pdf.add_page()
    _meta_bar(pdf)

    pdf.section_bar(f"収集したページ一覧（全 {len(scraped_pages)} ページ）")
    # 2列レイアウト（最大40ページ）
    col_pw = 86
    display_pages = scraped_pages[:40]
    _row_bgs = [(244, 241, 233), (238, 234, 222)]  # 交互背景
    for i, page in enumerate(display_pages):
        col_x = lm if i % 2 == 0 else lm + col_pw + 8
        if i % 2 == 0:
            if (297 - pdf.get_y() - 22) < 14:
                pdf.add_page()
                _meta_bar(pdf)
                pdf.set_font(pdf._font, "B", 8)
                pdf.set_text_color(*GRAY)
                pdf.cell(0, 5, "（続き）収集ページ一覧")
                pdf.ln(6)
                pdf.set_text_color(0, 0, 0)
            y_row_p = pdf.get_y()
        y_p2 = y_row_p
        title = page.get("title", "（タイトルなし）")[:22]
        url   = page.get("url", "")[:32]
        row_num = i // 2
        # カード背景（行ごとに交互）
        pdf.set_fill_color(*_row_bgs[row_num % 2])
        pdf.rect(col_x, y_p2, col_pw, 12, style="F")
        # 番号バッジ（ゴールド正方形）
        pdf.set_fill_color(*GOLD)
        pdf.rect(col_x, y_p2, 12, 12, style="F")
        pdf.set_xy(col_x, y_p2 + 1.5)
        pdf.set_font(pdf._font, "B", 7)
        pdf.set_text_color(*WHITE)
        pdf.cell(12, 9, str(i + 1), align="C")
        # タイトル
        pdf.set_xy(col_x + 14, y_p2 + 1.5)
        pdf.set_font(pdf._font, "B", 7.5)
        pdf.set_text_color(30, 24, 12)
        pdf.cell(col_pw - 15, 5, title)
        # URL
        pdf.set_xy(col_x + 14, y_p2 + 7)
        pdf.set_font(pdf._font, "", 6)
        pdf.set_text_color(*GRAY)
        pdf.cell(col_pw - 15, 4, url)
        if i % 2 == 1 or i == len(display_pages) - 1:
            pdf.set_y(y_p2 + 14)
    if len(scraped_pages) > 40:
        pdf.set_font(pdf._font, "", 7.5)
        pdf.set_text_color(*GRAY)
        pdf.set_x(lm)
        pdf.cell(0, 5, f"…他 {len(scraped_pages) - 40} ページ（合計 {len(scraped_pages)} ページ収集）")
        pdf.ln(5)
    pdf.set_text_color(0, 0, 0)

    # ─────────────────────────────────────────────
    # 最終ページ手前  ネクストアクション3選
    # ─────────────────────────────────────────────
    pdf.next_actions_page(site_report)

    # ─────────────────────────────────────────────
    # 最終ページ  各指標の見方ガイド
    # ─────────────────────────────────────────────
    pdf.guide_page()
    return bytes(pdf.output())


# ─────────────────────────────────────────────────────────────────
# 簡略版サマリーPDF（HTML+Playwright版・無料配布用）
# ─────────────────────────────────────────────────────────────────

def _build_summary_html(
    site_report: dict,
    site_url: str,
    profession: str,
    device_label: str,
) -> str:
    now = datetime.now().strftime("%Y\u5e74%m\u6708%d\u65e5")
    _device_clean = device_label.replace("\U0001f4bb", "").replace("\U0001f4f1", "").strip()

    def _b64img(path: str) -> str:
        if not os.path.exists(path):
            return ""
        with open(path, "rb") as fh:
            import base64 as _b64
            return _b64.b64encode(fh.read()).decode()

    char_b64 = _b64img(_CHARACTER_WORK_PATH) or _b64img(_CHARACTER_PATH)
    logo_html = ""
    header_logo_html = ""
    char_html = f'<img class="char-img" src="data:image/png;base64,{char_b64}">' if char_b64 else ""

    power_total, _, _, _ = _calc_power_score(site_report)
    power_pct = min(power_total, 100)

    rate = site_report.get("inquiry_rate", "-")
    try:
        rate_num = int("".join(c for c in rate if c.isdigit()))
    except Exception:
        rate_num = 0
    if rate_num >= 30:
        rate_color, rate_badge = "#16A34A", "\u9ad8\u6c34\u6e96"
    elif rate_num >= 12:
        rate_color, rate_badge = "#D97706", "\u6a19\u6e96"
    else:
        rate_color, rate_badge = "#B91C1C", "\u8981\u6539\u5584"

    mi = site_report.get("marketing_insights", {})
    drm = mi.get("drm_score", "-")
    drm_color = {"A": "#16A34A", "B": "#2563EB", "C": "#D97706", "D": "#B91C1C"}.get(drm, "#6B7280")
    drm_descs = {
        "A": "\u96c6\u5ba2\u30fb\u6559\u80b2\u30fb\u8ca9\u58f2\u304c\u5168\u6a5f\u80fd",
        "B": "1\uff5e2\u6539\u5584\u3067\u5927\u304d\u304f\u4f38\u3073\u308b",
        "C": "\u6539\u5584\u524d\u306b\u5e83\u544a\u306f\u5371\u967a",
        "D": "\u6839\u672c\u7684\u306a\u518d\u69cb\u7bc9\u304c\u5fc5\u8981",
    }
    drm_desc = drm_descs.get(drm, "")
    drm_pct  = {"A": 100, "B": 75, "C": 50, "D": 25}.get(str(drm), 0)

    def _esc(s: str) -> str:
        return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    strengths_html = ""
    for s in site_report.get("strengths", [])[:2]:
        strengths_html += (
            '<div class="sw-item s-item">' +
            f'<div class="sw-title s-title">{_esc(s.get("point",""))}</div>' +
            f'<div class="sw-body">{_esc(s.get("reason",""))}</div>' +
            '</div>'
        )

    weaknesses_html = ""
    for w in site_report.get("weaknesses", [])[:2]:
        weaknesses_html += (
            '<div class="sw-item w-item">' +
            f'<div class="sw-title w-title">{_esc(w.get("point",""))}</div>' +
            f'<div class="sw-body">{_esc(w.get("suggestion",""))}</div>' +
            '</div>'
        )

    impression = site_report.get("overall_impression", "")
    impression = (impression[:230] + "\u2026") if len(impression) > 230 else impression
    priority   = site_report.get("priority_action", "").strip()
    priority   = (priority[:230]   + "\u2026") if len(priority) > 230 else priority

    title    = f"{profession}\u3000\u7121\u6599\u30b5\u30a4\u30c8\u8a3a\u65ad\u30ec\u30dd\u30fc\u30c8" if profession else "\u7121\u6599\u30b5\u30a4\u30c8\u8a3a\u65ad\u30ec\u30dd\u30fc\u30c8"
    url_disp = (site_url[:80] + "\u2026") if len(site_url) > 80 else site_url

    css = """
@page { size: A4; margin: 0; }
* { box-sizing: border-box; margin: 0; padding: 0; }
body {
  font-family: 'Yu Gothic','Meiryo','Noto Sans CJK JP',sans-serif;
  background: #080808; color: #fff;
  width: 210mm; height: 297mm;
  display: flex; flex-direction: column;
  overflow: hidden; font-size: 10pt; line-height: 1.4;
}
.topline { height: 4px; background: #D4AF37; flex-shrink: 0; }
.header { background: #080808; padding: 16px 22px 12px; flex-shrink: 0; display: flex; align-items: flex-start; gap: 12px; }
.header-text { flex: 1; }
.header-logo { width: 42px; height: 42px; object-fit: contain; flex-shrink: 0; margin-top: 4px; }
.header-sub { font-size: 6.5pt; color: #D4AF37; letter-spacing: .1em; margin-bottom: 6px; }
.header-title { font-size: 19pt; font-weight: 700; color: #fff; margin-bottom: 6px; }
.header-meta { font-size: 7pt; color: #7a6a3a; margin-bottom: 3px; }
.header-disc { font-size: 5.5pt; color: #4a4020; }
.hd { height: 1px; background: #D4AF37; margin: 0 22px; }
.scores {
  display: flex; gap: 5px; padding: 12px 22px 18px;
  background: #f5f3ec; flex-shrink: 0;
}
.sc {
  flex: 1; background: #fff;
  border-top: 3px solid; padding: 12px 10px 10px;
  text-align: center; display: flex; flex-direction: column; align-items: center;
}
.sc-lbl { font-size: 6.5pt; font-weight: 700; margin-bottom: 6px; letter-spacing: .04em; }
.sc-num { font-size: 38pt; font-weight: 700; line-height: 1; margin-bottom: 3px; }
.sc-unit { font-size: 7pt; color: #6a5428; margin-bottom: 7px; }
.bar-bg { width: calc(100% - 20px); height: 4px; background: #ddd8c8; border-radius: 2px; margin-bottom: 8px; }
.bar-fill { height: 4px; border-radius: 2px; }
.badge {
  display: inline-block; padding: 3px 16px; border-radius: 3px;
  font-size: 8.5pt; font-weight: 700; color: #fff; margin-bottom: 7px;
}
.sc-guide { font-size: 5.5pt; color: #4a3a18; line-height: 1.8; }
.sc-desc { font-size: 7.5pt; color: #1a1510; font-weight: 700; margin-bottom: 5px; }
.sc-sub { font-size: 5.5pt; color: #4a3a18; line-height: 1.8; }
.content {
  flex: 1; background: #f5f3ec; padding: 12px 22px 10px;
  color: #1a1510; display: flex; flex-direction: column; gap: 10px; min-height: 0;
}
.impression { background: #fff; border-left: 4px solid #D4AF37; padding: 8px 12px; flex-shrink: 0; }
.sec-lbl { font-size: 6.5pt; font-weight: 700; color: #8a6a10; margin-bottom: 3px; }
.imp-text { font-size: 8pt; color: #1a1510; line-height: 1.65; }
.sw-row { display: flex; gap: 5px; flex: 1; min-height: 0; max-height: 68mm; }
.sw-col { flex: 1; display: flex; flex-direction: column; }
.sw-head { padding: 5px 10px; font-size: 8pt; font-weight: 700; color: #fff; flex-shrink: 0; }
.sw-items { flex: 1; display: flex; flex-direction: column; gap: 4px; }
.sw-item { background: #fff; border-left: 3px solid; padding: 6px 10px; flex: 1; overflow: hidden; }
.s-item { border-left-color: #166534; }
.w-item { border-left-color: #991b1b; }
.sw-title { font-size: 8pt; font-weight: 700; margin-bottom: 3px; }
.s-title { color: #166534; }
.w-title { color: #991b1b; }
.sw-body { font-size: 6.5pt; line-height: 1.6; color: #2a2a2a; }
.priority { background: #fff9e0; border-left: 4px solid #92400e; padding: 8px 12px; flex-shrink: 0; }
.pr-lbl { font-size: 6.5pt; font-weight: 700; color: #92400e; margin-bottom: 3px; }
.pr-text { font-size: 9pt; font-weight: 700; color: #2a1a02; line-height: 1.55; }
.guide { background: #e5e1d0; padding: 8px 12px; flex-shrink: 0; margin-top: auto; }
.guide-title { font-size: 8pt; font-weight: 700; color: #5a4a10; margin-bottom: 5px; }
.gr { display: flex; gap: 6px; margin-bottom: 3px; }
.gr-key { font-size: 6pt; font-weight: 700; color: #5a4a10; min-width: 96px; }
.gr-val { font-size: 6pt; color: #3a2a08; line-height: 1.6; flex: 1; }
.footer {
  background: #080808; border-top: 3px solid #D4AF37;
  flex-shrink: 0; position: relative;
  display: flex; align-items: flex-end;
  overflow: hidden; height: 62mm;
}
.char-img {
  position: absolute; bottom: 0; left: 0;
  height: 70mm; width: auto; object-fit: contain; opacity: 0.95;
  mix-blend-mode: screen;
  -webkit-mask-image: linear-gradient(to right, transparent 0%, black 42%);
  mask-image: linear-gradient(to right, transparent 0%, black 42%);
}
.footer-body {
  margin-left: 76mm; padding: 14px 22px 14px 0;
  display: flex; flex-direction: column; justify-content: center;
  flex: 1; align-self: center;
}
.logo-img { width: 22px; height: 22px; object-fit: contain; margin-bottom: 5px; }
.header-logo { width: 38px; height: 38px; object-fit: contain; flex-shrink: 0; margin-top: 2px; }
.ft-name { font-size: 11pt; font-weight: 700; color: #fff; margin-bottom: 2px; }
.ft-email { font-size: 7pt; color: #D4AF37; margin-bottom: 8px; }
.ft-div { height: 1px; background: #2a2010; margin-bottom: 8px; }
.ft-up-title { font-size: 8pt; font-weight: 700; color: #D4AF37; margin-bottom: 3px; }
.ft-up-text { font-size: 6.5pt; color: #8a7a4a; line-height: 1.65; }
"""

    return f"""<!DOCTYPE html>
<html lang="ja"><head><meta charset="UTF-8">
<style>{css}</style>
</head>
<body>
<div class="topline"></div>
<div class="header">
  <div class="header-text">
    <div class="header-sub">LIFE DESIGN LAB  —  トドク VCT  無料サイト診断レポート</div>
    <div class="header-title">{_esc(title)}</div>
    <div class="header-meta">{_esc(url_disp)}  |  {_esc(_device_clean)}  |  {now}</div>
    <div class="header-disc">本レポートはAIシンセティックデータ（付想顧客）による推定です。実際の数値は異なる場合があります。</div>
  </div>
  {header_logo_html}
</div>
<div class="hd"></div>

<div class="scores">
  <div class="sc" style="border-top-color:#D4AF37;">
    <div class="sc-lbl" style="color:#D4AF37;">AI 総合パワースコア</div>
    <div class="sc-num" style="color:#D4AF37;">{power_total}</div>
    <div class="sc-unit">/ 100点</div>
    <div class="bar-bg"><div class="bar-fill" style="width:{power_pct}%;background:#D4AF37;"></div></div>
    <div class="sc-guide">DRM + BrandZ + GEO の総合評価<br>（詳細内訳は詳細版レポートで）</div>
  </div>
  <div class="sc" style="border-top-color:{rate_color};">
    <div class="sc-lbl" style="color:{rate_color};">推定問い合わせ率</div>
    <div class="sc-num" style="color:{rate_color};">{_esc(rate)}</div>
    <span class="badge" style="background:{rate_color};">{rate_badge}</span>
    <div class="bar-bg"><div class="bar-fill" style="width:{min(rate_num,100)}%;background:{rate_color};"></div></div>
    <div class="sc-guide">30%以上=高水準<br>12～29%=標準　11%以下=要改善<br>付想顧客が回遊後に<br>問い合わせたいと感じた割合</div>
  </div>
  <div class="sc" style="border-top-color:{drm_color};">
    <div class="sc-lbl" style="color:{drm_color};">マーケティング評価 DRM</div>
    <div class="sc-num" style="color:{drm_color};">{_esc(drm)}</div>
    <div class="sc-desc">{_esc(drm_desc)}</div>
    <div class="bar-bg"><div class="bar-fill" style="width:{drm_pct}%;background:{drm_color};"></div></div>
    <div class="sc-sub">A=広告効果が出やすい<br>B=伸びしろあり<br>C=改善が先決<br>D=再構築が必要</div>
  </div>
</div>

<div class="content">
  <div class="impression">
    <div class="sec-lbl">第一印象  —  AIが付想顧客として見たサイト全体の総評</div>
    <div class="imp-text">{_esc(impression)}</div>
  </div>
  <div class="sw-row">
    <div class="sw-col">
      <div class="sw-head" style="background:#166534;">強み  —  集客に効いている点</div>
      <div class="sw-items">{strengths_html}</div>
    </div>
    <div class="sw-col">
      <div class="sw-head" style="background:#991b1b;">改善点  —  優先的に直す箇所</div>
      <div class="sw-items">{weaknesses_html}</div>
    </div>
  </div>
  <div class="priority">
    <div class="pr-lbl">今すぐやるべき最重要改善  —  これだけで問い合わせ数が最も変わる</div>
    <div class="pr-text">{_esc(priority)}</div>
  </div>
  <div class="guide">
    <div class="guide-title">このレポートの見方</div>
    <div class="gr"><span class="gr-key">- AI総合パワースコア</span><span class="gr-val">DRM・BrandZ・GEO 3軸の総合点（100点満点）。詳細な内訳は詳細版レポートで確認できます。</span></div>
    <div class="gr"><span class="gr-key">- 推定問い合わせ率</span><span class="gr-val">付想顧客が回遊後に問い合わせたいと感じた割合。30%以上=高水準。</span></div>
    <div class="gr"><span class="gr-key">- DRM（A～D）</span><span class="gr-val">集客→教育→販売の導線評価。Cは改善なく広告をかけると費用が無駄になるサイン。</span></div>
  </div>
</div>

<div class="footer">
  {char_html}
  <div class="footer-body">
    {logo_html}
    <div class="ft-name">LIFE DESIGN LAB</div>
    <div class="ft-email">inquiry.lifedesignlab@gmail.com</div>
    <div class="ft-div"></div>
    <div class="ft-up-title">詳細レポートをご希望の方へ</div>
    <div class="ft-up-text">このサマリーは診断結果のごく一部です。詳細版では強み・改善点の全項目、導線分析・競合比較・AIが生成した改善コピーをご確認いただけます。お気軽にご相談ください。</div>
  </div>
</div>
</body></html>"""

_chromium_ready = False

def _ensure_chromium() -> None:
    """Chromiumがなければ自動インストール（Streamlit Cloud対応）。"""
    global _chromium_ready
    if _chromium_ready:
        return
    import subprocess, sys
    try:
        subprocess.run(
            [sys.executable, "-m", "playwright", "install", "chromium"],
            capture_output=True, timeout=180,
        )
    except Exception:
        pass
    _chromium_ready = True


def _render_html_to_pdf(html: str) -> bytes | None:
    """PlaywrightでHTMLをA4 PDFに変換。失敗時はNone。"""
    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        return None
    _ensure_chromium()
    try:
        with sync_playwright() as p:
            kwargs = {"args": ["--no-sandbox", "--disable-dev-shm-usage", "--disable-gpu"]}
            browser = None
            for exe in ["/usr/bin/chromium-browser", "/usr/bin/chromium", None]:
                try:
                    if exe and not os.path.exists(exe):
                        continue
                    kw = dict(kwargs)
                    if exe:
                        kw["executable_path"] = exe
                    browser = p.chromium.launch(**kw)
                    break
                except Exception:
                    continue
            if browser is None:
                return None
            try:
                page = browser.new_page()
                page.set_content(html, wait_until="domcontentloaded")
                return page.pdf(
                    format="A4",
                    print_background=True,
                    margin={"top": "0", "right": "0", "bottom": "0", "left": "0"},
                )
            finally:
                browser.close()
    except Exception:
        return None


def generate_summary_pdf(
    site_report: dict,
    site_url: str = "",
    profession: str = "",
    device_label: str = "PC",
) -> bytes:
    """診断結果の要点を1枚にまとめた簡略版PDF（HTML優先・fpdf2フォールバック）。"""
    html = _build_summary_html(site_report, site_url, profession, device_label)
    pdf = _render_html_to_pdf(html)
    if pdf:
        return pdf
    return _generate_summary_pdf_fpdf2(site_report, site_url, profession, device_label)


def _generate_summary_pdf_fpdf2(
    site_report: dict,
    site_url: str = "",
    profession: str = "",
    device_label: str = "PC",
) -> bytes:
    """fpdf2によるサマリーPDFフォールバック（Streamlit Cloud用）。"""
    font_r, font_b = _resolve_font()
    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.add_page()
    pdf.set_auto_page_break(auto=False)

    if font_r:
        pdf.add_font("F", "", font_r)
        pdf.add_font("F", "B", font_b or font_r)
        fn = "F"
    else:
        fn = "helvetica"

    W, lm = 210, 11
    inner = W - lm * 2
    now = datetime.now().strftime("%Y年%m月%d日")
    _device_clean = device_label.replace("💻", "").replace("📱", "").strip()

    # 背景
    pdf.set_fill_color(*NAVY_DARK)
    pdf.rect(0, 0, W, 108, style="F")
    pdf.set_fill_color(248, 246, 240)
    pdf.rect(0, 108, W, 189, style="F")
    pdf.set_fill_color(*GOLD)
    pdf.rect(0, 0, W, 1.8, style="F")

    # ヘッダー
    pdf.set_xy(lm, 8)
    pdf.set_font(fn, "B", 6)
    pdf.set_text_color(*GOLD)
    pdf.cell(inner, 4, "LIFE DESIGN LAB  -  トドク VCT  無料サイト診断レポート")
    pdf.set_xy(lm, 14)
    pdf.set_font(fn, "B", 14)
    pdf.set_text_color(*WHITE)
    title_text = f"{profession}  無料サイト診断レポート" if profession else "無料サイト診断レポート"
    pdf.cell(inner, 8, title_text)
    pdf.set_xy(lm, 24)
    pdf.set_font(fn, "", 6)
    pdf.set_text_color(160, 148, 110)
    url_disp = (site_url[:80] + "...") if len(site_url) > 80 else site_url
    pdf.cell(inner, 4, f"{url_disp}  |  {_device_clean}  |  {now}")
    pdf.set_xy(lm, 29)
    pdf.set_font(fn, "", 5.5)
    pdf.set_text_color(110, 102, 72)
    pdf.cell(inner, 3.5, "本レポートはAIシンセティックデータ（仮想顧客）による推定です。")

    # スコアパネル
    power_total, _, _, _ = _calc_power_score(site_report)
    rate = site_report.get("inquiry_rate", "-")
    try:
        rate_num = int("".join(c for c in rate if c.isdigit()))
    except Exception:
        rate_num = 0
    if rate_num >= 30:
        rate_color, rate_badge = GREEN_MID, "高水準"
    elif rate_num >= 12:
        rate_color, rate_badge = (217, 119, 6), "標準"
    else:
        rate_color, rate_badge = RED_T, "要改善"

    mi = site_report.get("marketing_insights", {})
    drm = mi.get("drm_score", "-")
    drm_color = {"A": GREEN_MID, "B": BLUE_MID, "C": (217,119,6), "D": RED_T}.get(drm, GRAY)

    panel_y, panel_h, gap = 38, 64, 3
    pw = (inner - gap * 2) / 3
    _bg = (24, 20, 10)

    for idx, (px, top_c, label, big_val, sub1, sub2) in enumerate([
        (lm,           GOLD,       "AI 総合パワースコア", str(power_total), "/ 100点",  "DRM+BrandZ+GEO（詳細は有料版で）"),
        (lm+pw+gap,    rate_color, "推定問い合わせ率",   rate,             rate_badge, "30%以上=高水準 12〜29%=標準"),
        (lm+pw*2+gap*2,drm_color, "マーケティング DRM", drm,              "",         {"A":"広告効果大","B":"伸びしろあり","C":"改善が先決","D":"再構築必要"}.get(drm,"")),
    ]):
        pdf.set_fill_color(*_bg)
        pdf.rect(px, panel_y, pw, panel_h, style="F")
        pdf.set_fill_color(*top_c)
        pdf.rect(px, panel_y, pw, 1.2, style="F")
        pdf.set_xy(px, panel_y + 4)
        pdf.set_font(fn, "B", 5.5)
        pdf.set_text_color(*top_c)
        pdf.cell(pw, 4, label, align="C")
        pdf.set_xy(px, panel_y + 9)
        pdf.set_font(fn, "B", 32)
        pdf.set_text_color(*top_c)
        pdf.cell(pw, 18, big_val, align="C")
        pdf.set_xy(px, panel_y + 28)
        pdf.set_font(fn, "B", 7)
        pdf.set_text_color(*WHITE)
        pdf.cell(pw, 5, sub1, align="C")
        pdf.set_xy(px, panel_y + 35)
        pdf.set_font(fn, "", 5.5)
        pdf.set_text_color(140, 130, 90)
        pdf.cell(pw, 4, sub2, align="C")

    # パワースコアバー
    bar_y = panel_y + 43
    pdf.set_fill_color(45, 40, 22)
    pdf.rect(lm + 6, bar_y, pw - 12, 3, style="F")
    fill_w = min(power_total / 100.0, 1.0) * (pw - 12)
    if fill_w > 0:
        pdf.set_fill_color(*GOLD)
        pdf.rect(lm + 6, bar_y, fill_w, 3, style="F")

    # 第一印象
    imp_y = 110
    impression = site_report.get("overall_impression", "")
    pdf.set_fill_color(*WHITE)
    pdf.rect(lm, imp_y, inner, 16, style="F")
    pdf.set_fill_color(*GOLD)
    pdf.rect(lm, imp_y, 2.5, 16, style="F")
    pdf.set_xy(lm + 5, imp_y + 2)
    pdf.set_font(fn, "B", 6)
    pdf.set_text_color(120, 100, 40)
    pdf.cell(inner - 7, 4, "第一印象  —  AIが仮想顧客として見た総評")
    pdf.set_xy(lm + 5, imp_y + 7)
    pdf.set_font(fn, "", 7)
    pdf.set_text_color(40, 35, 25)
    pdf.multi_cell(inner - 7, 4.5, (impression[:108] + "...") if len(impression) > 108 else impression)

    # 強み・改善点
    col_w = (inner - 3) / 2
    lx, rx = lm, lm + col_w + 3
    sec_y = 129
    pdf.set_fill_color(*GREEN_T)
    pdf.rect(lx, sec_y, col_w, 7, style="F")
    pdf.set_xy(lx, sec_y)
    pdf.set_font(fn, "B", 7)
    pdf.set_text_color(*WHITE)
    pdf.cell(col_w, 7, "  強み  —  集客に効いている点")
    pdf.set_fill_color(*RED_T)
    pdf.rect(rx, sec_y, col_w, 7, style="F")
    pdf.set_xy(rx, sec_y)
    pdf.cell(col_w, 7, "  改善点  —  優先的に直す箇所")

    for i in range(2):
        iy = sec_y + 8 + i * 26
        if i < len(site_report.get("strengths", [])):
            s = site_report["strengths"][i]
            pdf.set_fill_color(*WHITE)
            pdf.rect(lx, iy, col_w, 25, style="F")
            pdf.set_fill_color(*GREEN_T)
            pdf.rect(lx, iy, 2, 25, style="F")
            pdf.set_xy(lx + 4, iy + 2)
            pdf.set_font(fn, "B", 7)
            pdf.set_text_color(*GREEN_T)
            pdf.multi_cell(col_w - 6, 4, s.get("point", "")[:35])
            pdf.set_xy(lx + 4, iy + 10)
            pdf.set_font(fn, "", 6)
            pdf.set_text_color(60, 80, 60)
            pdf.multi_cell(col_w - 6, 3.8, s.get("reason", "")[:65])
        if i < len(site_report.get("weaknesses", [])):
            w = site_report["weaknesses"][i]
            pdf.set_fill_color(*WHITE)
            pdf.rect(rx, iy, col_w, 25, style="F")
            pdf.set_fill_color(*RED_T)
            pdf.rect(rx, iy, 2, 25, style="F")
            pdf.set_xy(rx + 4, iy + 2)
            pdf.set_font(fn, "B", 7)
            pdf.set_text_color(*RED_T)
            pdf.multi_cell(col_w - 6, 4, w.get("point", "")[:35])
            pdf.set_xy(rx + 4, iy + 10)
            pdf.set_font(fn, "", 6)
            pdf.set_text_color(100, 50, 50)
            cur = w.get("current_text", "")
            sug = w.get("suggestion", "")
            combined = (f"現在: {cur[:30]}\n改善後: {sug[:30]}" if cur else f"改善後: {sug[:55]}")
            pdf.multi_cell(col_w - 6, 3.8, combined)

    # 最重要改善
    pdf.set_fill_color(255, 252, 230)
    pdf.rect(lm, 186, inner, 22, style="F")
    pdf.set_fill_color(*AMBER_T)
    pdf.rect(lm, 186, 3, 22, style="F")
    pdf.set_xy(lm + 6, 188)
    pdf.set_font(fn, "B", 6)
    pdf.set_text_color(*AMBER_T)
    pdf.cell(inner - 8, 4, "今すぐやるべき最重要改善")
    priority = site_report.get("priority_action", "").strip()
    if priority:
        pdf.set_xy(lm + 6, 193)
        pdf.set_font(fn, "B", 7.5)
        pdf.set_text_color(50, 30, 5)
        pdf.multi_cell(inner - 8, 4.5, priority[:130])

    # 用語解説
    pdf.set_fill_color(228, 224, 212)
    pdf.rect(lm, 210, inner, 19, style="F")
    pdf.set_xy(lm + 3, 212)
    pdf.set_font(fn, "B", 7)
    pdf.set_text_color(80, 65, 25)
    pdf.cell(inner - 6, 5, "このレポートの見方")
    for j, (lbl, dsc) in enumerate([
        ("AI総合パワースコア", "DRM・BrandZ・GEO 3軸の総合点（100点満点）。詳細は詳細版レポートで。"),
        ("推定問い合わせ率",   "仮想顧客が回遊後に問い合わせたいと感じた割合。30%以上=高水準。"),
        ("DRM（A〜D）",       "集客→教育→販売の導線評価。Cは改善なく広告をかけると費用が無駄に。"),
    ]):
        gy = 218 + j * 4
        pdf.set_xy(lm + 3, gy)
        pdf.set_font(fn, "B", 5.5)
        pdf.set_text_color(80, 65, 25)
        pdf.cell(36, 3.5, f"- {lbl}")
        pdf.set_xy(lm + 39, gy)
        pdf.set_font(fn, "", 5.5)
        pdf.set_text_color(60, 48, 18)
        pdf.cell(inner - 41, 3.5, dsc)

    # フッター
    pdf.set_fill_color(*NAVY_DARK)
    pdf.rect(0, 232, W, 65, style="F")
    pdf.set_fill_color(*GOLD)
    pdf.rect(0, 232, W, 1.2, style="F")
    if os.path.exists(_LOGO_PATH):
        pdf.image(_LOGO_PATH, x=lm, y=236, h=11)
    pdf.set_xy(lm + 15, 236)
    pdf.set_font(fn, "B", 8.5)
    pdf.set_text_color(*WHITE)
    pdf.cell(inner - 15, 6, "LIFE DESIGN LAB")
    pdf.set_xy(lm + 15, 243)
    pdf.set_font(fn, "", 6.5)
    pdf.set_text_color(*GOLD)
    pdf.cell(inner - 15, 4, "inquiry.lifedesignlab@gmail.com")
    pdf.set_fill_color(40, 35, 16)
    pdf.rect(lm, 250, inner, 0.3, style="F")
    pdf.set_xy(lm, 253)
    pdf.set_font(fn, "B", 7)
    pdf.set_text_color(*GOLD)
    pdf.cell(inner, 5, "詳細レポートをご希望の方へ")
    pdf.set_xy(lm, 259)
    pdf.set_font(fn, "", 6.5)
    pdf.set_text_color(200, 190, 160)
    pdf.cell(inner, 5, "このサマリーは診断結果のごく一部です。詳細版では全項目・導線分析・AIコピーをご確認いただけます。")
    pdf.set_xy(lm, 265)
    pdf.cell(inner, 5, "お気軽にご相談ください。")

    return bytes(pdf.output())


def generate_script_pdf(script: dict, site_url: str, profession: str) -> bytes:
    """30分Zoom解説台本をDOCX形式で返す（Googleドキュメント用）。"""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    doc = Document()

    # ページ余白
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    def _heading(text: str, level: int = 1):
        p = doc.add_heading(text, level=level)
        p.runs[0].font.color.rgb = RGBColor(0x08, 0x08, 0x08)
        return p

    def _label(text: str):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xB4, 0x5A, 0x09)
        return p

    def _script(text: str):
        p = doc.add_paragraph()
        run = p.add_run(f"▶ {text}")
        run.font.size = Pt(10.5)
        p.paragraph_format.left_indent = Cm(0.5)
        return p

    def _normal(text: str, italic: bool = False, color: RGBColor = None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(10)
        if italic:
            run.italic = True
        if color:
            run.font.color.rgb = color
        p.paragraph_format.left_indent = Cm(0.5)
        return p

    def _before(text: str):
        p = doc.add_paragraph()
        run = p.add_run(f"現在: {text}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)
        p.paragraph_format.left_indent = Cm(1)
        return p

    def _after(text: str):
        p = doc.add_paragraph()
        run = p.add_run(f"改善後: {text}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x15, 0x80, 0x3D)
        run.bold = True
        p.paragraph_format.left_indent = Cm(1)
        return p

    def _qa(q: str, a: str, idx: int):
        p = doc.add_paragraph()
        run = p.add_run(f"Q{idx}. {q}")
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0x78, 0x50, 0x00)
        p2 = doc.add_paragraph()
        run2 = p2.add_run(f"A. {a}")
        run2.font.size = Pt(10)
        p2.paragraph_format.left_indent = Cm(0.5)
        doc.add_paragraph()

    # ── タイトル ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("【コンサルタント専用】30分 Zoom 解説台本")
    run.bold = True
    run.font.size = Pt(16)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run(f"{profession}  |  {site_url}")
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x80, 0x70, 0x40)
    doc.add_paragraph()

    # ── ① オープニング ──
    opening = script.get("opening", {})
    _label(f"① オープニング　{opening.get('duration','2分')}")
    _script(opening.get("script", ""))
    doc.add_paragraph()

    # ── ② スコア概要 ──
    score = script.get("score_overview", {})
    _label(f"② スコア概要　{score.get('duration','5分')}")
    _script(score.get("script", ""))
    for tp in score.get("talking_points", []):
        _normal(f"・{tp}")
    doc.add_paragraph()

    # ── ③ 強み ──
    strengths = script.get("strengths_section", {})
    _label(f"③ 強み　{strengths.get('duration','5分')}")
    _script(strengths.get("script", ""))
    for item in strengths.get("items", []):
        _normal(f"✓ {item.get('point','')}", color=RGBColor(0x15, 0x80, 0x3D))
        _normal(item.get("explanation", ""), italic=True)
    doc.add_paragraph()

    # ── ④ 改善点 Before/After ──
    ba = script.get("before_after_section", {})
    _label(f"④ 改善点 Before / After　{ba.get('duration','10分')}")
    _script(ba.get("script", ""))
    doc.add_paragraph()
    for item in ba.get("items", []):
        p = doc.add_paragraph()
        run = p.add_run(f"✗ {item.get('point','')}")
        run.bold = True
        run.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)
        if item.get("current_text"):
            _before(item["current_text"])
        if item.get("suggestion"):
            _after(item["suggestion"])
        if item.get("explanation"):
            _normal(f"解説: {item['explanation']}", italic=True, color=RGBColor(0x60, 0x50, 0x20))
        doc.add_paragraph()

    # ── ⑤ GEOスコア ──
    geo = script.get("geo_section", {})
    if geo.get("script"):
        _label(f"⑤ GEOスコア解説　{geo.get('duration','3分')}")
        _script(geo.get("script", ""))
        for item in geo.get("items", []):
            p = doc.add_paragraph()
            run = p.add_run(f"{item.get('category','')}  {item.get('score','')}")
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x78, 0x50, 0x00)
            _normal(item.get("explanation", ""), italic=True)
        doc.add_paragraph()

    # ── ⑥ 導線・ナビゲーション ──
    nav = script.get("navigation_section", {})
    if nav.get("script") and nav.get("items"):
        _label(f"⑥ 導線・ナビゲーション分析　{nav.get('duration','3分')}")
        _script(nav.get("script", ""))
        for item in nav.get("items", []):
            p = doc.add_paragraph()
            run = p.add_run(f"【{item.get('page','')}】{item.get('issue','')}")
            run.bold = True
            run.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)
            _normal(f"改善: {item.get('suggestion','')}", color=RGBColor(0x15, 0x80, 0x3D))
        doc.add_paragraph()

    # ── ⑦ 最重要アクション ──
    pri = script.get("priority_action", {})
    _label(f"⑦ 最重要アクション　{pri.get('duration','5分')}")
    _script(pri.get("script", ""))
    if pri.get("action"):
        _normal(f"アクション: {pri['action']}", color=RGBColor(0x92, 0x40, 0x0E))
    if pri.get("expected_result"):
        _normal(f"期待される変化: {pri['expected_result']}", color=RGBColor(0x15, 0x80, 0x3D))
    doc.add_paragraph()

    # ── ⑧ 次のステップ ──
    nxt = script.get("next_steps", {})
    _label(f"⑧ 次のステップ　{nxt.get('duration','2分')}")
    _script(nxt.get("script", ""))
    if nxt.get("kindle_mention"):
        _normal(f"（次の提案として）{nxt['kindle_mention']}", color=RGBColor(0x6B, 0x46, 0x1D))
    doc.add_paragraph()

    # ── ⑨ 予想Q&A ──
    qa_list = script.get("qa_list", [])
    _label(f"⑨ 予想 Q&A　{len(qa_list)}問")
    doc.add_paragraph()
    for i, qa in enumerate(qa_list, 1):
        _qa(qa.get("question", ""), qa.get("answer", ""), i)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
